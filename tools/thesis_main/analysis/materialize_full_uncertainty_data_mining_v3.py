"""Comprehensive advisor-facing annotation data audit, version 3.

This materializer extends the reviewed v2 package without replacing it.  It
keeps all computable records, reports every orthogonal exclusion reason, adds a
pair-only sensitivity lane for tasks with two annotators, enumerates all
Manual/Semi convergence and expansion cases, audits proposal anchoring, and
separates metric-decreasing edits from claims about real-world harm.

Terminology contract
--------------------
A negative GT/utility change is called ``negative_metric_change``.  It is not
interpreted as a harmful edit because the operational GT and the Manhattan
representation can themselves be incomplete or force a particular fit.
Likewise, an ``acceptable`` proposal tag followed by a corner edit is not by
itself strong annotation uncertainty: the analysis distinguishes micro/local
adjustment, material coordinate change, topology change, mode change and GT
alignment.
"""
from __future__ import annotations

import argparse
import ast
import hashlib
import json
import math
import os
import re
import shutil
import statistics
import textwrap
import zipfile
from collections import Counter, defaultdict
from dataclasses import dataclass
from datetime import datetime, timezone
from itertools import combinations
from pathlib import Path
from typing import Any, Iterable, Mapping, Sequence

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt
from openpyxl import Workbook, load_workbook
from openpyxl.drawing.image import Image as XLImage
from openpyxl.formatting.rule import ColorScaleRule
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter
from scipy import stats

ROOT = Path(__file__).resolve().parents[3]
V2 = ROOT / "analysis_results" / "full_uncertainty_data_mining_20260821_v2"
C1_AUDIT = ROOT / "analysis_results" / "c1_formal_audit_20260802_v16_final" / "c1_formal_audit_20260802_7fcacc5c2d6c_bf5def46_6bc67c03"
C1_MINING = ROOT / "analysis_results" / "annotation_uncertainty_manual_semi_20260820_v2"
PACKAGE = ROOT / "analysis_results" / "paper_a_data_mining_package_20260820_v1" / "curated"
PERSISTENT = ROOT / "analysis_results" / "persistent_disagreement_diagnostic_20260819_v1"
DEFAULT_OUTPUT = ROOT / "analysis_results" / "full_uncertainty_data_mining_20260821_v3"
SEED = 20260821
PANORAMA_WIDTH = 1024.0
PANORAMA_HEIGHT = 512.0

MISSING = {"", "none", "null", "nan", "na", "n/a", "not_evaluable", "not_identifiable"}
IMAGE_SUFFIXES = {".jpg", ".jpeg", ".png", ".webp", ".bmp", ".tif", ".tiff"}


def clean(value: Any) -> str:
    if value is None:
        return ""
    text = str(value).strip().lstrip("\ufeff").strip()
    return "" if text.lower() in MISSING else text


def truth(value: Any) -> bool:
    return clean(value).lower() in {"1", "true", "yes", "y", "passed", "valid", "eligible", "matched"}


def worker_id(value: Any) -> str:
    token = clean(value).upper()
    if token.startswith("W"):
        token = token[1:]
    return str(int(token)) if token.isdigit() else token


def normalise_stage(value: Any) -> str:
    token = clean(value)
    aliases = {
        "C2A-RP-B1": "C2-A-RP-B1",
        "C2A-RP-B2": "C2-A-RP-B2",
        "C2A_RP_B1": "C2-A-RP-B1",
        "C2A_RP_B2": "C2-A-RP-B2",
    }
    return aliases.get(token, token)


def normalise_condition(value: Any) -> str:
    token = clean(value).lower().replace("-", "_")
    if "semi" in token or "assist" in token or "model" in token:
        return "semi"
    if "manual" in token:
        return "manual"
    if token == "oos":
        return "oos"
    return token or "unknown"


def numeric(series: pd.Series) -> pd.Series:
    return pd.to_numeric(series, errors="coerce")


def read_csv(path: Path, required: bool = True) -> pd.DataFrame:
    if not path.is_file():
        if required:
            raise FileNotFoundError(path)
        return pd.DataFrame()
    return pd.read_csv(path, encoding="utf-8-sig", low_memory=False)


def first_existing(paths: Sequence[Path], required: bool = True) -> Path | None:
    for path in paths:
        if path.is_file():
            return path
    if required:
        raise FileNotFoundError("none of: " + ", ".join(map(str, paths)))
    return None


def read_first(paths: Sequence[Path], required: bool = True) -> pd.DataFrame:
    path = first_existing(paths, required=required)
    return read_csv(path, required=required) if path else pd.DataFrame()


def write_csv(path: Path, frame: pd.DataFrame) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    frame.to_csv(path, index=False, encoding="utf-8", lineterminator="\n")


def write_json(path: Path, value: Any) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(value, ensure_ascii=False, indent=2, sort_keys=True, default=json_default) + "\n", encoding="utf-8")


def json_default(value: Any) -> Any:
    if isinstance(value, (np.integer, np.floating, np.bool_)):
        return value.item()
    if isinstance(value, Path):
        return value.as_posix()
    if pd.isna(value):
        return None
    raise TypeError(type(value).__name__)


def sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for block in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest()


def parse_jsonish(value: Any) -> Any:
    if isinstance(value, (dict, list, tuple)):
        return value
    text = clean(value)
    if not text:
        return None
    for loader in (json.loads, ast.literal_eval):
        try:
            return loader(text)
        except Exception:
            pass
    return None


def parse_points(value: Any) -> list[list[float]]:
    payload = parse_jsonish(value)
    if payload is None:
        return []
    if isinstance(payload, dict):
        for key in ("corners_px", "ordered_geometry", "canonical_geometry", "geometry", "points"):
            if key in payload:
                points = parse_points(payload[key])
                if points:
                    return points
        return []
    if isinstance(payload, (list, tuple)):
        if payload and all(isinstance(item, (list, tuple)) and len(item) >= 2 for item in payload):
            result: list[list[float]] = []
            for item in payload:
                try:
                    x, y = float(item[0]), float(item[1])
                except Exception:
                    return []
                if not (math.isfinite(x) and math.isfinite(y)):
                    return []
                result.append([x, y])
            return result
        for item in payload:
            points = parse_points(item)
            if points:
                return points
    return []


def strip_image_name(value: Any) -> str:
    text = clean(value).replace("\\", "/").rsplit("/", 1)[-1]
    suffix = Path(text).suffix.lower()
    return text[: -len(suffix)] if suffix in IMAGE_SUFFIXES else text


def recursive_strings(value: Any) -> Iterable[str]:
    if isinstance(value, dict):
        for child in value.values():
            yield from recursive_strings(child)
    elif isinstance(value, (list, tuple)):
        for child in value:
            yield from recursive_strings(child)
    elif isinstance(value, str):
        yield value


def extract_image_reference(value: Any) -> str:
    payload = parse_jsonish(value)
    for text in recursive_strings(payload):
        if Path(text.split("?", 1)[0]).suffix.lower() in IMAGE_SUFFIXES:
            return text
    return ""


def entropy_from_counts(counts: Iterable[int | float]) -> float:
    arr = np.asarray([float(value) for value in counts if float(value) > 0], dtype=float)
    if not len(arr):
        return float("nan")
    p = arr / arr.sum()
    return float(-(p * np.log(p)).sum())


def safe_spearman(x: pd.Series, y: pd.Series) -> tuple[float | None, float | None, int]:
    pair = pd.DataFrame({"x": numeric(x), "y": numeric(y)}).dropna()
    if len(pair) < 3 or pair["x"].nunique() < 2 or pair["y"].nunique() < 2:
        return None, None, len(pair)
    result = stats.spearmanr(pair["x"], pair["y"])
    return float(result.statistic), float(result.pvalue), len(pair)


def quantile_band(series: pd.Series, value: float | None) -> str:
    values = numeric(series).dropna()
    if value is None or not math.isfinite(float(value)) or values.empty:
        return "not_computable"
    positive = values[values > 0]
    if float(value) <= 0:
        return "none"
    if len(positive) < 4:
        return "positive_edit"
    q25, q75 = positive.quantile([0.25, 0.75])
    if value <= q25:
        return "micro_or_small"
    if value <= q75:
        return "moderate"
    return "large"


def circular_dx(left: float, right: float, width: float = PANORAMA_WIDTH) -> float:
    delta = abs(left - right) % width
    return min(delta, width - delta)


def geometry_signature(points: list[list[float]]) -> dict[str, Any]:
    if len(points) < 4 or len(points) % 2:
        return {
            "n_points": len(points), "n_pairs": np.nan, "top_y": np.nan,
            "bottom_y": np.nan, "room_height": np.nan, "wall_x_json": "[]",
            "seam_pair_count": np.nan, "valid": False,
        }
    pairs = []
    for index in range(0, len(points), 2):
        a, b = points[index], points[index + 1]
        top, bottom = sorted((float(a[1]), float(b[1])))
        x = (float(a[0]) + float(b[0])) / 2.0
        pairs.append((x % PANORAMA_WIDTH, top, bottom))
    xs = sorted(item[0] for item in pairs)
    top = float(np.median([item[1] for item in pairs]))
    bottom = float(np.median([item[2] for item in pairs]))
    seam_count = sum(x < 64 or x > PANORAMA_WIDTH - 64 for x in xs)
    return {
        "n_points": len(points), "n_pairs": len(pairs), "top_y": top,
        "bottom_y": bottom, "room_height": bottom - top,
        "wall_x_json": json.dumps(xs), "seam_pair_count": seam_count, "valid": True,
    }


def wall_x_distance(left_json: str, right_json: str) -> float | None:
    left, right = parse_jsonish(left_json) or [], parse_jsonish(right_json) or []
    if not left or len(left) != len(right):
        return None
    a, b = np.asarray(left, dtype=float), np.asarray(right, dtype=float)
    candidates = []
    for shift in range(len(b)):
        candidate = np.roll(b, shift)
        distances = [circular_dx(float(x), float(y)) for x, y in zip(a, candidate)]
        candidates.append(float(np.sqrt(np.mean(np.square(distances)))))
    return min(candidates) if candidates else None


def building_from_task(value: Any) -> str:
    task = strip_image_name(value)
    return task.split("_", 1)[0] if "_" in task else "not_identifiable"


def load_data() -> dict[str, pd.DataFrame]:
    data = {
        "unified": read_csv(V2 / "UNIFIED_SUBMISSION_EVIDENCE_REVIEWED.csv"),
        "raw": read_csv(PACKAGE / "raw_annotation_fact.csv"),
        "meta_long": read_csv(V2 / "RAW_META_LABEL_RESPONSE_LONG.csv"),
        "meta_task": read_csv(V2 / "META_LABEL_TASK_UNCERTAINTY.csv"),
        "geometry_task": read_csv(V2 / "GEOMETRY_TASK_UNCERTAINTY_ALL_STAGES.csv"),
        "geometry_pair": read_csv(V2 / "GEOMETRY_PAIRWISE_DISPERSION_ALL_STAGES.csv"),
        "crowd_gt": read_csv(V2 / "C1_CROWD_GT_CONFLICT_TASKS.csv"),
        "crowd_cluster": read_csv(V2 / "C1_CROWD_GT_CLUSTER_METRICS.csv"),
        "membership": read_csv(V2 / "C1_WORKER_TASK_MODE_MEMBERSHIP.csv"),
        "worker_view": read_csv(V2 / "WORKER_VIEWPOINT_STABILITY.csv"),
        "time": read_csv(V2 / "TIME_SUBMISSION_EVIDENCE_REVIEWED.csv"),
        "population_tasks": read_csv(C1_MINING / "POPULATION_TASK_METRICS.csv"),
        "population_summary": read_csv(C1_MINING / "POPULATION_SENSITIVITY.csv"),
        "row_inclusion": read_csv(C1_MINING / "ROW_INCLUSION_CLASSIFICATION.csv"),
        "quality_context": read_first([
            C1_MINING / "QUALITY_DATA_MINING_CONTEXTS.csv",
            C1_AUDIT / "c1_gt_quality_analysis.csv",
        ]),
        "quality_task": read_first([
            C1_MINING / "QUALITY_DATA_MINING_TASK_METRICS.csv",
            C1_MINING / "QUALITY_AUXILIARY.csv",
        ]),
        "excluded_peer": read_csv(C1_MINING / "EXCLUDED_WORKER_PEER_COMPARISONS.csv", required=False),
        "excluded_impact": read_csv(C1_MINING / "EXCLUDED_CONTEXT_IMPACT.csv", required=False),
        "semi_review": read_csv(PACKAGE / "semi_review_fact.csv"),
        "persistent": read_csv(PERSISTENT / "PERSISTENT_DISAGREEMENT_TASKS.csv"),
    }
    for frame in data.values():
        if "stage" in frame:
            frame["stage"] = frame["stage"].map(normalise_stage)
        if "condition" in frame:
            frame["condition"] = frame["condition"].map(normalise_condition)
        if "worker_id" in frame:
            frame["worker_id"] = frame["worker_id"].map(worker_id)
        for name in ("worker_id_left", "worker_id_right", "excluded_worker_id", "formal_peer_worker_id"):
            if name in frame:
                frame[name] = frame[name].map(worker_id)
        if "base_task_id" in frame:
            frame["base_task_id"] = frame["base_task_id"].map(strip_image_name)
    raw = data["raw"]
    if "image_reference" not in raw:
        raw["image_reference"] = raw.get("task_data_json", pd.Series([""] * len(raw))).map(extract_image_reference)
    if "image_reference" in raw:
        raw["image_reference"] = raw["image_reference"].map(clean)
    return data


def build_image_reference_map(raw: pd.DataFrame) -> dict[str, str]:
    result: dict[str, str] = {}
    if raw.empty:
        return result
    for row in raw.itertuples(index=False):
        task = clean(getattr(row, "base_task_id", ""))
        ref = clean(getattr(row, "image_reference", ""))
        if task and ref and task not in result:
            result[task] = ref
    return result


def parse_secondary_flags(value: Any) -> set[str]:
    payload = parse_jsonish(value)
    if isinstance(payload, list):
        return {clean(item) for item in payload if clean(item)}
    text = clean(value)
    return {token.strip() for token in re.split(r"[;|]", text) if token.strip()} if text else set()


def orthogonal_record_flags(unified: pd.DataFrame) -> tuple[pd.DataFrame, pd.DataFrame, pd.DataFrame, pd.DataFrame]:
    frame = unified.copy()
    for column in ("active_time_observed_seconds", "lead_time_seconds", "iou_to_gt", "n_corners"):
        if column not in frame:
            frame[column] = np.nan
        frame[column] = numeric(frame[column])
    if "geometry_computable" not in frame:
        frame["geometry_computable"] = False
    frame["geometry_computable_bool"] = frame["geometry_computable"].map(truth)
    frame["active_time_computable_bool"] = frame["active_time_observed_seconds"].notna()
    frame["lead_time_computable_bool"] = frame["lead_time_seconds"].notna()
    frame["quality_computable_bool"] = frame["iou_to_gt"].notna()
    frame["formal_use_allowed_bool"] = frame.get("formal_use_allowed", pd.Series([False] * len(frame))).map(truth)
    parsed = frame.get("secondary_exclusion_flags", pd.Series([""] * len(frame))).map(parse_secondary_flags)
    process = frame.get("worker_process_class", pd.Series([""] * len(frame))).astype(str).str.lower()
    assignment = frame.get("assignment_provenance", pd.Series([""] * len(frame))).astype(str).str.lower()
    primary = frame.get("primary_exclusion_class", pd.Series([""] * len(frame))).astype(str).str.lower()
    scope = frame.get("task_final_scope", pd.Series([""] * len(frame))).astype(str).str.lower()
    status = frame.get("worker_status_class", pd.Series([""] * len(frame))).astype(str).str.lower()

    definitions: list[tuple[str, str, pd.Series]] = [
        ("administrative_exclusion", "行政退出/行政排除", process.str.contains("administr") | status.str.contains("administr") | parsed.map(lambda s: any("administr" in x for x in s))),
        ("outside_assignment", "任务分配之外的提交", assignment.str.contains("outside") | process.str.contains("outside") | parsed.map(lambda s: any("outside_assignment" in x for x in s))),
        ("formal_assignment_ineligible", "旧正式分配资格不足", ~frame["formal_use_allowed_bool"] | parsed.map(lambda s: any("formal_assignment:not_eligible" in x for x in s))),
        ("independence_ineligible", "独立性条件不足", parsed.map(lambda s: any(x.startswith("independence:") for x in s))),
        ("scope_oos_or_ineligible", "范围外或 Scope 不满足", scope.eq("oos") | primary.str.contains("oos_scope") | parsed.map(lambda s: any(x.startswith("scope:") for x in s))),
        ("geometry_not_computable", "几何不可计算", ~frame["geometry_computable_bool"] | parsed.map(lambda s: any(x.startswith("geometry:") for x in s))),
        ("quality_not_computable", "GT/质量指标不可计算", ~frame["quality_computable_bool"] | parsed.map(lambda s: any(x.startswith("quality:") for x in s))),
        ("active_time_not_computable", "Active time 不可计算", ~frame["active_time_computable_bool"] | parsed.map(lambda s: any("active_time:not_evaluable" in x for x in s))),
        ("manual_only_no_semi_candidate", "只有 Manual、没有 Semi 候选", primary.str.contains("manual_only") | frame.get("has_semi_candidate", pd.Series([False] * len(frame))).map(lambda x: not truth(x))),
        ("historical_retained_worker", "历史状态工人但仍保留", status.str.contains("historical")),
        ("noncanonical_or_raw_only_role", "非 canonical 或仅 raw 可审计记录", ~frame.get("record_role", pd.Series(["stage_canonical"] * len(frame))).astype(str).eq("stage_canonical")),
    ]
    for code, _label, mask in definitions:
        frame[f"flag__{code}"] = mask.fillna(False).astype(bool)
    flag_cols = [f"flag__{code}" for code, _, _ in definitions]
    frame["orthogonal_flag_count"] = frame[flag_cols].sum(axis=1)
    frame["orthogonal_flags"] = frame.apply(
        lambda row: ";".join(code for code, _, _ in definitions if bool(row[f"flag__{code}"])), axis=1
    )

    summary_rows = []
    outcome_rows = []
    for code, label, _mask in definitions:
        subset = frame[frame[f"flag__{code}"]]
        complement = frame[~frame[f"flag__{code}"]]
        summary_rows.append({
            "reason_code": code, "reason_zh": label,
            "record_count": len(subset), "worker_count": subset["worker_id"].nunique(),
            "task_count": subset["base_task_id"].nunique(), "stage_count": subset["stage"].nunique(),
            "geometry_computable_count": int(subset["geometry_computable_bool"].sum()),
            "quality_computable_count": int(subset["quality_computable_bool"].sum()),
            "active_time_computable_count": int(subset["active_time_computable_bool"].sum()),
            "lead_time_computable_count": int(subset["lead_time_computable_bool"].sum()),
            "formal_use_allowed_count": int(subset["formal_use_allowed_bool"].sum()),
            "record_share": len(subset) / len(frame) if len(frame) else np.nan,
            "note": "orthogonal reason; rows may occur in multiple reasons",
        })
        for population, group in (("reason_present", subset), ("reason_absent", complement)):
            outcome_rows.append({
                "reason_code": code, "reason_zh": label, "population": population,
                "record_count": len(group), "worker_count": group["worker_id"].nunique(),
                "task_count": group["base_task_id"].nunique(),
                "geometry_computable_rate": float(group["geometry_computable_bool"].mean()) if len(group) else np.nan,
                "quality_computable_rate": float(group["quality_computable_bool"].mean()) if len(group) else np.nan,
                "iou_to_gt_mean": float(group["iou_to_gt"].mean()) if group["iou_to_gt"].notna().any() else np.nan,
                "iou_to_gt_median": float(group["iou_to_gt"].median()) if group["iou_to_gt"].notna().any() else np.nan,
                "active_time_median_seconds": float(group["active_time_observed_seconds"].median()) if group["active_time_observed_seconds"].notna().any() else np.nan,
                "lead_time_median_seconds": float(group["lead_time_seconds"].median()) if group["lead_time_seconds"].notna().any() else np.nan,
                "n_corners_median": float(group["n_corners"].median()) if group["n_corners"].notna().any() else np.nan,
                "interpretation": "descriptive subgroup; overlapping reasons and non-random composition",
            })
    summary = pd.DataFrame(summary_rows).sort_values("record_count", ascending=False)
    outcomes = pd.DataFrame(outcome_rows)
    overlap_rows = []
    for left, right in combinations([(code, label) for code, label, _ in definitions], 2):
        left_code, left_label = left
        right_code, right_label = right
        both = frame[f"flag__{left_code}"] & frame[f"flag__{right_code}"]
        overlap_rows.append({
            "reason_left": left_code, "reason_left_zh": left_label,
            "reason_right": right_code, "reason_right_zh": right_label,
            "overlap_record_count": int(both.sum()),
            "jaccard": float(both.sum() / (frame[f"flag__{left_code}"] | frame[f"flag__{right_code}"]).sum())
            if (frame[f"flag__{left_code}"] | frame[f"flag__{right_code}"]).sum() else np.nan,
        })
    overlap = pd.DataFrame(overlap_rows).sort_values("overlap_record_count", ascending=False)
    return frame, summary, outcomes, overlap


def raw_out_of_task(raw: pd.DataFrame, unified: pd.DataFrame) -> tuple[pd.DataFrame, pd.DataFrame]:
    frame = raw.copy()
    for column in ("stage", "condition", "base_task_id", "worker_id", "annotation_id", "canonical_annotation_id", "canonical_join_status"):
        if column not in frame:
            frame[column] = ""
    frame["stage"] = frame["stage"].map(normalise_stage)
    frame["condition"] = frame["condition"].map(normalise_condition)
    frame["worker_id"] = frame["worker_id"].map(worker_id)
    known_by_stage = unified.groupby("stage")["base_task_id"].apply(lambda s: set(map(clean, s))).to_dict()
    canonical_ids = set(map(clean, unified.get("annotation_id", pd.Series(dtype=str)))) | set(map(clean, unified.get("canonical_annotation_id", pd.Series(dtype=str))))
    outside_keys = set(
        zip(
            unified.loc[unified.get("flag__outside_assignment", pd.Series([False] * len(unified))).astype(bool), "stage"],
            unified.loc[unified.get("flag__outside_assignment", pd.Series([False] * len(unified))).astype(bool), "base_task_id"],
            unified.loc[unified.get("flag__outside_assignment", pd.Series([False] * len(unified))).astype(bool), "worker_id"],
        )
    )
    frame["base_task_not_in_selected_stage"] = [clean(task) not in known_by_stage.get(stage, set()) for stage, task in zip(frame["stage"], frame["base_task_id"])]
    frame["annotation_id_not_selected"] = ~frame["annotation_id"].map(clean).isin(canonical_ids)
    frame["canonical_join_nonmatched"] = ~frame["canonical_join_status"].astype(str).str.lower().isin({"matched", "canonical", "canonical_joined", "selected"})
    frame["outside_assignment_key"] = [(stage, task, worker) in outside_keys for stage, task, worker in zip(frame["stage"], frame["base_task_id"], frame["worker_id"])]
    frame["out_of_task_or_unplanned"] = frame[["base_task_not_in_selected_stage", "outside_assignment_key"]].any(axis=1)
    frame["raw_revision_or_nonselected"] = frame[["annotation_id_not_selected", "canonical_join_nonmatched"]].any(axis=1)
    frame["audit_class"] = frame.apply(
        lambda row: ";".join(
            code for code, value in (
                ("base_task_not_in_selected_stage", row["base_task_not_in_selected_stage"]),
                ("outside_assignment", row["outside_assignment_key"]),
                ("annotation_version_not_selected", row["annotation_id_not_selected"]),
                ("canonical_join_nonmatched", row["canonical_join_nonmatched"]),
            ) if bool(value)
        ) or "selected_or_matched",
        axis=1,
    )
    selected = frame[frame["out_of_task_or_unplanned"] | frame["raw_revision_or_nonselected"]].copy()
    summary = selected.groupby(["stage", "condition", "audit_class"], dropna=False).agg(
        raw_row_count=("annotation_id", "size"), worker_count=("worker_id", "nunique"),
        task_count=("base_task_id", "nunique"),
        lead_time_observed_count=("lead_time_seconds", lambda values: numeric(values).notna().sum()) if "lead_time_seconds" in selected else ("annotation_id", "size"),
    ).reset_index() if not selected.empty else pd.DataFrame()
    return selected, summary


def semi_convergence_expansion(data: Mapping[str, pd.DataFrame], image_map: Mapping[str, str]) -> pd.DataFrame:
    population = data["population_tasks"].copy()
    population["threshold"] = numeric(population["threshold"])
    target = population[
        population["population"].eq("all_canonical_planned")
        & np.isclose(population["threshold"], 0.95)
    ].copy()
    if len(target) != 25:
        raise AssertionError(f"expected 25 paired Manual/Semi tasks, observed {len(target)}")
    for column in [name for name in target.columns if name.startswith(("manual_", "semi_", "delta_"))]:
        target[column] = numeric(target[column])
    quality = data["quality_task"].copy()
    quality_columns = [column for column in (
        "base_task_id", "delta_all_computable_iou", "delta_iou_to_gt",
        "manual_all_computable_iou_mean", "semi_all_computable_iou_mean",
        "manual_iou_to_gt_mean", "semi_iou_to_gt_mean",
    ) if column in quality]
    if quality_columns:
        quality = quality[quality_columns].drop_duplicates("base_task_id")
        for column in quality_columns:
            if column != "base_task_id":
                quality[column] = numeric(quality[column])
        target = target.merge(quality, how="left", on="base_task_id")
    target["delta_quality_iou"] = target.get("delta_all_computable_iou", target.get("delta_iou_to_gt", np.nan))
    target["uncertainty_direction"] = np.select(
        [target["delta_shannon_entropy"].lt(0), target["delta_shannon_entropy"].gt(0)],
        ["semi_convergence", "semi_expansion"], default="no_entropy_change",
    )
    abs_delta = target["delta_shannon_entropy"].abs()
    q1, q2 = abs_delta.quantile([1 / 3, 2 / 3])
    target["magnitude_band"] = pd.cut(
        abs_delta, bins=[-np.inf, q1, q2, np.inf], labels=["lower_third", "middle_third", "upper_third"], include_lowest=True
    ).astype(str)
    target["mode_transition"] = target.apply(
        lambda row: f"{row.get('manual_mode_count', np.nan):g}->{row.get('semi_mode_count', np.nan):g}"
        if pd.notna(row.get("manual_mode_count")) and pd.notna(row.get("semi_mode_count")) else "not_computable", axis=1
    )
    target["supported_multimodality_transition"] = target.apply(
        lambda row: f"{int(row.get('manual_supported_multimodality', 0))}->{int(row.get('semi_supported_multimodality', 0))}"
        if pd.notna(row.get("manual_supported_multimodality")) and pd.notna(row.get("semi_supported_multimodality")) else "not_computable", axis=1
    )
    target["quality_direction"] = np.select(
        [target["delta_quality_iou"].gt(0.01), target["delta_quality_iou"].lt(-0.01)],
        ["positive_metric_change", "negative_metric_change"], default="near_zero_or_not_computable",
    )
    target["image_reference"] = target["base_task_id"].map(image_map).fillna("")
    target["interpretation_zh"] = target.apply(
        lambda row: (
            "Semi 的工人间几何分布更集中；是否更正确需单独查看 GT、proposal 与编辑。"
            if row["uncertainty_direction"] == "semi_convergence" else
            "Semi 的工人间几何分布更分散；可能来自 proposal 歧义、不同修改方向或图像本身多解。"
            if row["uncertainty_direction"] == "semi_expansion" else
            "Manual 与 Semi 的任务级熵相同或无法区分。"
        ), axis=1
    )
    return target.sort_values("delta_shannon_entropy").reset_index(drop=True)


def prepare_semi_review(data: Mapping[str, pd.DataFrame]) -> pd.DataFrame:
    review = data["semi_review"].copy()
    for column in ("U_initial", "U_final", "delta_U", "geometry_edit_rmse_px", "geometry_edit_rmse_panorama_diagonal_normalized"):
        if column not in review:
            review[column] = np.nan
        review[column] = numeric(review[column])
    if "edited_binary" in review:
        review["edited_bool"] = review["edited_binary"].map(truth)
    else:
        review["edited_bool"] = review["geometry_edit_rmse_px"].fillna(0).gt(0)
    review["worker_id"] = review.get("worker_id", pd.Series([""] * len(review))).map(worker_id)
    review["condition"] = "semi"
    review["stage"] = review.get("stage", pd.Series([""] * len(review))).map(normalise_stage)
    review["base_task_id"] = review["base_task_id"].map(strip_image_name)
    review["metric_change_class"] = np.select(
        [review["delta_U"].gt(0.01), review["delta_U"].lt(-0.01), review["delta_U"].notna()],
        ["positive_metric_change", "negative_metric_change", "near_zero_metric_change"],
        default="metric_not_computable",
    )
    edit_basis = review["geometry_edit_rmse_px"].where(review["geometry_edit_rmse_px"].notna(), review["geometry_edit_rmse_panorama_diagonal_normalized"] * math.hypot(PANORAMA_WIDTH, PANORAMA_HEIGHT))
    review["edit_magnitude_px_equivalent"] = edit_basis
    review["edit_magnitude_band"] = [quantile_band(edit_basis, value) for value in edit_basis]
    initial_hash = review.get("initial_geometry_hash", review.get("initialization_geometry_hash", pd.Series([""] * len(review)))).map(clean)
    final_hash = review.get("final_geometry_hash", review.get("geometry_hash", pd.Series([""] * len(review)))).map(clean)
    review["geometry_hash_changed"] = initial_hash.ne("") & final_hash.ne("") & initial_hash.ne(final_hash)
    review["topology_changed"] = False
    for left_name, right_name in (
        ("initial_point_count", "final_point_count"),
        ("initial_n_corners", "final_n_corners"),
        ("initial_pair_count", "final_pair_count"),
    ):
        if left_name in review and right_name in review:
            review["topology_changed"] |= numeric(review[left_name]).ne(numeric(review[right_name])) & numeric(review[left_name]).notna() & numeric(review[right_name]).notna()
    return review


def proposal_stage_summary(review: pd.DataFrame) -> pd.DataFrame:
    return review.groupby(["stage", "condition"], dropna=False).agg(
        row_count=("base_task_id", "size"), task_count=("base_task_id", "nunique"),
        worker_count=("worker_id", "nunique"), edited_count=("edited_bool", "sum"),
        edited_rate=("edited_bool", "mean"), positive_metric_change_count=("metric_change_class", lambda values: sum(value == "positive_metric_change" for value in values)),
        negative_metric_change_count=("metric_change_class", lambda values: sum(value == "negative_metric_change" for value in values)),
        near_zero_metric_change_count=("metric_change_class", lambda values: sum(value == "near_zero_metric_change" for value in values)),
        metric_not_computable_count=("metric_change_class", lambda values: sum(value == "metric_not_computable" for value in values)),
        edit_rmse_median=("edit_magnitude_px_equivalent", "median"),
        initial_metric_mean=("U_initial", "mean"), final_metric_mean=("U_final", "mean"), delta_metric_mean=("delta_U", "mean"),
    ).reset_index()


def proposal_task_analysis(review: pd.DataFrame, convergence: pd.DataFrame, image_map: Mapping[str, str]) -> pd.DataFrame:
    c1 = review[(review["stage"].eq("C1")) & review["base_task_id"].isin(set(convergence["base_task_id"]))].copy()
    task = c1.groupby("base_task_id", as_index=False).agg(
        semi_review_row_count=("worker_id", "size"), worker_count=("worker_id", "nunique"),
        edit_count=("edited_bool", "sum"), edit_rate=("edited_bool", "mean"),
        edit_rmse_median=("edit_magnitude_px_equivalent", "median"), edit_rmse_mean=("edit_magnitude_px_equivalent", "mean"),
        positive_metric_change_count=("metric_change_class", lambda values: sum(value == "positive_metric_change" for value in values)),
        negative_metric_change_count=("metric_change_class", lambda values: sum(value == "negative_metric_change" for value in values)),
        near_zero_metric_change_count=("metric_change_class", lambda values: sum(value == "near_zero_metric_change" for value in values)),
        initial_metric_mean=("U_initial", "mean"), final_metric_mean=("U_final", "mean"), delta_metric_mean=("delta_U", "mean"),
        topology_change_count=("topology_changed", "sum"), geometry_hash_change_count=("geometry_hash_changed", "sum"),
    )
    columns = [
        "base_task_id", "building_id", "delta_shannon_entropy", "manual_shannon_entropy", "semi_shannon_entropy",
        "manual_mode_count", "semi_mode_count", "manual_largest_mode_share", "semi_largest_mode_share",
        "delta_quality_iou", "uncertainty_direction", "magnitude_band", "image_reference",
    ]
    task = convergence[[column for column in columns if column in convergence]].merge(task, how="left", on="base_task_id")
    task["edit_rate"] = numeric(task["edit_rate"])
    task["edit_rmse_median"] = numeric(task["edit_rmse_median"])
    positive_edits = numeric(c1["edit_magnitude_px_equivalent"]).dropna()
    q25 = positive_edits[positive_edits > 0].quantile(0.25) if (positive_edits > 0).any() else np.nan
    q75 = positive_edits[positive_edits > 0].quantile(0.75) if (positive_edits > 0).any() else np.nan

    def classify(row: pd.Series) -> str:
        direction = row.get("uncertainty_direction")
        rate = row.get("edit_rate")
        median = row.get("edit_rmse_median")
        quality = row.get("delta_quality_iou")
        if direction == "semi_convergence" and pd.notna(rate) and rate <= 0.25:
            return "convergence_with_high_proposal_retention"
        if direction == "semi_convergence" and pd.notna(median) and pd.notna(q25) and median <= q25:
            return "convergence_after_small_shared_adjustments"
        if direction == "semi_convergence" and pd.notna(median) and pd.notna(q75) and median > q75:
            return "convergence_despite_substantial_revisions"
        if direction == "semi_convergence" and pd.notna(quality) and quality < -0.01:
            return "convergence_with_negative_GT_metric_tradeoff"
        if direction == "semi_convergence":
            return "convergence_with_mixed_revision_evidence"
        if direction == "semi_expansion" and pd.notna(rate) and rate > 0.5:
            return "proposal_followed_by_divergent_revisions"
        if direction == "semi_expansion":
            return "semi_expansion_with_limited_or_mixed_revisions"
        return "no_entropy_change"

    task["proposal_response_pattern"] = task.apply(classify, axis=1)
    task["anchor_evidence"] = task.apply(
        lambda row: (
            "consistent_with_passive_anchor_retention"
            if row["proposal_response_pattern"] == "convergence_with_high_proposal_retention" else
            "consistent_with_anchor_plus_shared_adjustment"
            if row["proposal_response_pattern"] in {"convergence_after_small_shared_adjustments", "convergence_with_mixed_revision_evidence"} else
            "not_clear_or_counterevidence"
        ), axis=1
    )
    task["structure_discovery_evidence"] = task.apply(
        lambda row: (
            "consistent_with_but_does_not_prove_shared_structure_discovery"
            if row.get("positive_metric_change_count", 0) > row.get("negative_metric_change_count", 0)
            and row.get("edit_count", 0) > 0
            and row.get("uncertainty_direction") != "semi_expansion" else
            "not_supported_by_available_GT_metric"
        ), axis=1
    )
    task["forced_fit_tradeoff_evidence"] = task.apply(
        lambda row: (
            "possible_metric_or_representation_tradeoff"
            if row.get("uncertainty_direction") == "semi_convergence"
            and ((pd.notna(row.get("delta_quality_iou")) and row.get("delta_quality_iou") < -0.01)
                 or row.get("negative_metric_change_count", 0) > row.get("positive_metric_change_count", 0))
            else "not_identified"
        ), axis=1
    )
    task["interpretation_boundary_zh"] = "模式集中只说明工人最终几何更接近；不能单独证明 proposal 正确、工人发现了结构或发生了不良修改。"
    return task.sort_values("delta_shannon_entropy").reset_index(drop=True)


def aggregate_worker_tags(meta_long: pd.DataFrame) -> pd.DataFrame:
    if meta_long.empty:
        return pd.DataFrame(columns=["stage", "condition", "base_task_id", "worker_id"])
    frame = meta_long.copy()
    frame["stage"] = frame["stage"].map(normalise_stage)
    frame["condition"] = frame["condition"].map(normalise_condition)
    frame["worker_id"] = frame["worker_id"].map(worker_id)
    frame["choice_code"] = frame["choice_code"].map(clean)
    grouped = frame.groupby(["stage", "condition", "base_task_id", "worker_id"], dropna=False).agg(
        tag_codes=("choice_code", lambda values: ";".join(sorted(set(filter(None, map(clean, values)))))),
        tag_labels_zh=("choice_label_zh", lambda values: ";".join(sorted(set(filter(None, map(clean, values)))))),
        acceptable_tag=("choice_code", lambda values: any(clean(value) == "acceptable" for value in values)),
        trivial_tag=("choice_code", lambda values: any(clean(value) == "trivial" for value in values)),
        in_scope_tag=("choice_code", lambda values: any(clean(value) == "in_scope" for value in values)),
        model_issue_tag=("meta_group", lambda values: any(clean(value) == "model_issue" for value in values)),
        explicit_model_problem_tag=("choice_code", lambda values: any(clean(value) not in {"", "acceptable"} for value in values)),
    ).reset_index()
    return grouped


def tag_behavior_analysis(data: Mapping[str, pd.DataFrame], review: pd.DataFrame, image_map: Mapping[str, str]) -> tuple[pd.DataFrame, pd.DataFrame, pd.DataFrame]:
    tags = aggregate_worker_tags(data["meta_long"])
    semi = review[review["condition"].eq("semi")].copy()
    row = semi.merge(tags, how="left", on=["stage", "condition", "base_task_id", "worker_id"])
    for column in ("acceptable_tag", "trivial_tag", "in_scope_tag", "model_issue_tag", "explicit_model_problem_tag"):
        row[column] = row.get(column, False).fillna(False).astype(bool)
    quality = data["quality_context"].copy()
    if not quality.empty:
        quality["worker_id"] = quality["worker_id"].map(worker_id)
        quality["condition"] = quality["condition"].map(normalise_condition)
        quality["base_task_id"] = quality["base_task_id"].map(strip_image_name)
        keep = [column for column in (
            "base_task_id", "condition", "worker_id", "iou_to_gt", "worker_caused_structural_failure",
            "gt_score_computable", "task_final_scope", "geometry_reference_status",
        ) if column in quality]
        quality = quality[keep].drop_duplicates(["base_task_id", "condition", "worker_id"])
        row = row.merge(quality, how="left", on=["base_task_id", "condition", "worker_id"])
    if "iou_to_gt" not in row:
        row["iou_to_gt"] = np.nan
    row["iou_to_gt"] = numeric(row["iou_to_gt"])
    row["structural_failure_bool"] = row.get("worker_caused_structural_failure", pd.Series([False] * len(row))).map(truth)
    membership = data["membership"].copy()
    if not membership.empty:
        membership["worker_id"] = membership["worker_id"].map(worker_id)
        keep = [column for column in (
            "base_task_id", "condition", "worker_id", "cluster_rank", "cluster_support", "cluster_count",
            "is_largest_mode", "is_supported_minority_mode", "n_pairs", "task_centered_n_pairs",
        ) if column in membership]
        membership = membership[keep].drop_duplicates(["base_task_id", "condition", "worker_id"])
        row = row.merge(membership, how="left", on=["base_task_id", "condition", "worker_id"])
    time = data["time"].copy()
    if not time.empty:
        time["worker_id"] = time["worker_id"].map(worker_id)
        keep = [column for column in (
            "stage", "condition", "base_task_id", "worker_id", "active_time_observed_seconds", "lead_time_seconds",
            "active_time_measurement_class", "active_time_source", "lead_time_source",
        ) if column in time]
        time = time[keep].drop_duplicates(["stage", "condition", "base_task_id", "worker_id"])
        row = row.merge(time, how="left", on=["stage", "condition", "base_task_id", "worker_id"])
    row["active_time_observed_seconds"] = numeric(row.get("active_time_observed_seconds", pd.Series([np.nan] * len(row))))
    row["lead_time_seconds"] = numeric(row.get("lead_time_seconds", pd.Series([np.nan] * len(row))))
    stage_q90 = row.groupby(["stage", "condition"])["active_time_observed_seconds"].transform(lambda values: numeric(values).quantile(0.90))
    row["long_active_time"] = row["active_time_observed_seconds"].gt(stage_q90) & stage_q90.notna()
    row["low_gt_alignment_absolute"] = row["iou_to_gt"].lt(0.80)
    task_median = row.groupby("base_task_id")["iou_to_gt"].transform("median")
    row["low_gt_alignment_relative"] = row["iou_to_gt"].notna() & task_median.notna() & row["iou_to_gt"].lt(task_median - 0.10)
    row["nonlargest_mode"] = row.get("is_largest_mode", pd.Series([False] * len(row))).notna() & ~row.get("is_largest_mode", pd.Series([False] * len(row))).map(truth)
    row["material_edit"] = row["edit_magnitude_band"].isin({"moderate", "large"}) | row["topology_changed"]
    row["image_reference"] = row["base_task_id"].map(image_map).fillna("")

    def codes(record: pd.Series) -> list[str]:
        result: list[str] = []
        if record["acceptable_tag"] and record["edited_bool"]:
            result.append("acceptable_proposal_with_any_corner_edit")
        if record["acceptable_tag"] and record["material_edit"]:
            result.append("acceptable_proposal_with_material_or_topology_edit")
        if record["acceptable_tag"] and record["metric_change_class"] == "negative_metric_change":
            result.append("acceptable_proposal_with_negative_metric_change")
        if record["acceptable_tag"] and record["nonlargest_mode"]:
            result.append("acceptable_proposal_but_final_nonlargest_mode")
        if record["acceptable_tag"] and (record["low_gt_alignment_absolute"] or record["low_gt_alignment_relative"]):
            result.append("acceptable_proposal_but_low_GT_alignment")
        if record["trivial_tag"] and (record["low_gt_alignment_absolute"] or record["low_gt_alignment_relative"]):
            result.append("trivial_tag_but_low_GT_alignment")
        if record["trivial_tag"] and record["nonlargest_mode"]:
            result.append("trivial_tag_but_final_nonlargest_mode")
        if record["trivial_tag"] and record["structural_failure_bool"]:
            result.append("trivial_tag_but_structural_failure_flag")
        if record["trivial_tag"] and record["long_active_time"]:
            result.append("trivial_tag_but_long_active_time")
        if record["trivial_tag"] and record["material_edit"]:
            result.append("trivial_tag_but_material_edit")
        if record["explicit_model_problem_tag"] and not record["edited_bool"]:
            result.append("model_problem_tag_without_geometry_edit")
        if record["in_scope_tag"] and clean(record.get("task_final_scope")) == "oos":
            result.append("in_scope_tag_on_final_OOS_task")
        return result

    row["case_codes"] = row.apply(lambda record: ";".join(codes(record)), axis=1)
    row["has_case"] = row["case_codes"].ne("")

    def interpretation(record: pd.Series) -> str:
        if record["acceptable_tag"] and record["edited_bool"]:
            if record["edit_magnitude_band"] == "micro_or_small" and not record["topology_changed"] and not record["nonlargest_mode"] and not record["low_gt_alignment_absolute"]:
                return "可接受标签后的小幅同模式调整，更符合局部精度修正；单独不足以说明强不确定性。"
            return "标签与几何行为不完全一致；较大坐标/拓扑变化或模式偏离构成残余标注不确定性的证据。"
        if record["trivial_tag"] and (record["low_gt_alignment_absolute"] or record["nonlargest_mode"] or record["structural_failure_bool"]):
            return "主观难度判断与几何准确性/共识结构不一致；简单标签不能替代质量检验。"
        if record["explicit_model_problem_tag"] and not record["edited_bool"]:
            return "工人识别了问题但未改变几何；问题可能不在角点、可能被界面限制，或标签与操作脱节。"
        return "无预定义标签—行为不一致。"

    row["uncertainty_interpretation_zh"] = row.apply(interpretation, axis=1)
    cases = row[row["has_case"]].copy().sort_values(["case_codes", "base_task_id", "worker_id"])
    exploded_rows = []
    for record in cases.to_dict("records"):
        for code in record["case_codes"].split(";"):
            exploded_rows.append({"case_code": code, **record})
    exploded = pd.DataFrame(exploded_rows)
    summary = exploded.groupby(["stage", "case_code"], dropna=False).agg(
        row_count=("base_task_id", "size"), task_count=("base_task_id", "nunique"),
        worker_count=("worker_id", "nunique"), median_iou=("iou_to_gt", "median"),
        median_edit_px=("edit_magnitude_px_equivalent", "median"),
        nonlargest_mode_rate=("nonlargest_mode", "mean"),
        negative_metric_change_rate=("metric_change_class", lambda values: np.mean([value == "negative_metric_change" for value in values])),
    ).reset_index() if not exploded.empty else pd.DataFrame()
    return row, cases, summary


def dual_annotator_sensitivity(data: Mapping[str, pd.DataFrame], image_map: Mapping[str, str]) -> tuple[pd.DataFrame, pd.DataFrame]:
    tasks = data["geometry_task"].copy()
    tasks["observed_worker_count"] = numeric(tasks["observed_worker_count"])
    tasks["geometry_computable_worker_count"] = numeric(tasks["geometry_computable_worker_count"])
    selected = tasks[(tasks["observed_worker_count"].eq(2)) & (tasks["geometry_computable_worker_count"].eq(2))].copy()
    pairs = data["geometry_pair"].copy()
    if "cyclic_rmse_diagonal_normalized" in pairs:
        pairs["cyclic_rmse_diagonal_normalized"] = numeric(pairs["cyclic_rmse_diagonal_normalized"])
    pair_agg = pairs.groupby(["stage", "condition", "base_task_id"], as_index=False).agg(
        pair_row_count=("worker_id_left", "size"), worker_id_left=("worker_id_left", "first"),
        worker_id_right=("worker_id_right", "first"), same_topology=("same_topology", lambda values: all(map(truth, values))),
        cyclic_rmse_diagonal_normalized=("cyclic_rmse_diagonal_normalized", "median"),
        n_corners_left=("n_corners_left", "first"), n_corners_right=("n_corners_right", "first"),
    ) if not pairs.empty else pd.DataFrame()
    if not pair_agg.empty:
        selected = selected.merge(pair_agg, how="left", on=["stage", "condition", "base_task_id"])
    selected["image_reference"] = selected["base_task_id"].map(image_map).fillna("")
    selected["topology_agreement_class"] = np.where(selected.get("same_topology", False).map(truth), "same_topology", "different_topology")
    rmse = numeric(selected.get("cyclic_rmse_diagonal_normalized", pd.Series([np.nan] * len(selected))))
    if rmse.notna().sum() >= 4:
        q1, q2 = rmse.dropna().quantile([1 / 3, 2 / 3])
        selected["coordinate_disagreement_band"] = pd.cut(rmse, [-np.inf, q1, q2, np.inf], labels=["lower_third", "middle_third", "upper_third"]).astype(str)
    else:
        selected["coordinate_disagreement_band"] = "not_evaluable"
    selected["sensitivity_role"] = "two_annotator_pairwise_sensitivity_without_GT_majority_claim"
    summary = selected.groupby(["stage", "condition"], dropna=False).agg(
        two_annotator_task_count=("base_task_id", "nunique"),
        same_topology_count=("topology_agreement_class", lambda values: sum(value == "same_topology" for value in values)),
        different_topology_count=("topology_agreement_class", lambda values: sum(value == "different_topology" for value in values)),
        cyclic_rmse_median=("cyclic_rmse_diagonal_normalized", "median"),
        cyclic_rmse_q25=("cyclic_rmse_diagonal_normalized", lambda values: numeric(values).quantile(0.25)),
        cyclic_rmse_q75=("cyclic_rmse_diagonal_normalized", lambda values: numeric(values).quantile(0.75)),
    ).reset_index()
    return selected.sort_values(["stage", "condition", "base_task_id"]), summary


def load_geometry_map() -> dict[str, list[list[float]]]:
    path = C1_AUDIT / "c1_canonical_geometry.jsonl"
    result: dict[str, list[list[float]]] = {}
    if not path.is_file():
        return result
    with path.open(encoding="utf-8-sig") as handle:
        for line in handle:
            if not line.strip():
                continue
            item = json.loads(line)
            identity = clean(item.get("canonical_annotation_id") or item.get("annotation_id"))
            points = parse_points(item.get("corners_px") or item.get("geometry") or item)
            if identity and points:
                result[identity] = points
    return result


def cluster_geometry_summaries(cluster: pd.DataFrame, geometry_map: Mapping[str, list[list[float]]]) -> pd.DataFrame:
    rows = []
    for record in cluster.to_dict("records"):
        identities = [clean(value) for value in clean(record.get("annotation_ids")).split(";") if clean(value)]
        signatures = [geometry_signature(geometry_map[identity]) for identity in identities if identity in geometry_map]
        valid = [item for item in signatures if item["valid"]]
        n_pairs = [item["n_pairs"] for item in valid]
        top = [item["top_y"] for item in valid]
        bottom = [item["bottom_y"] for item in valid]
        height = [item["room_height"] for item in valid]
        seam = [item["seam_pair_count"] for item in valid]
        wall_signatures = [parse_jsonish(item["wall_x_json"]) or [] for item in valid]
        mode_pairs = int(pd.Series(n_pairs).mode().iloc[0]) if n_pairs else np.nan
        matching = [values for values in wall_signatures if len(values) == mode_pairs] if pd.notna(mode_pairs) else []
        wall_median: list[float] = []
        if matching:
            reference = np.asarray(sorted(matching[0]), dtype=float)
            aligned = [reference]
            for values in matching[1:]:
                candidate = np.asarray(sorted(values), dtype=float)
                options = [np.roll(candidate, shift) for shift in range(len(candidate))]
                best = min(options, key=lambda option: np.mean([circular_dx(x, y) ** 2 for x, y in zip(reference, option)]))
                aligned.append(best)
            wall_median = list(np.median(np.vstack(aligned), axis=0))
        rows.append({
            **record,
            "geometry_computable_member_count": len(valid),
            "cluster_n_pairs_median": float(np.median(n_pairs)) if n_pairs else np.nan,
            "cluster_n_pairs_values": ";".join(map(lambda x: f"{x:g}", n_pairs)),
            "cluster_top_y_median": float(np.median(top)) if top else np.nan,
            "cluster_bottom_y_median": float(np.median(bottom)) if bottom else np.nan,
            "cluster_room_height_median": float(np.median(height)) if height else np.nan,
            "cluster_seam_pair_count_median": float(np.median(seam)) if seam else np.nan,
            "cluster_wall_x_median_json": json.dumps(wall_median),
        })
    return pd.DataFrame(rows)


def crowd_gt_geometric_causes(data: Mapping[str, pd.DataFrame], image_map: Mapping[str, str]) -> tuple[pd.DataFrame, pd.DataFrame]:
    conflict = data["crowd_gt"].copy()
    if len(conflict) != 101:
        raise AssertionError(f"expected 101 crowd-GT task-condition records, observed {len(conflict)}")
    cluster = cluster_geometry_summaries(data["crowd_cluster"], load_geometry_map())
    meta = data["meta_long"].copy()
    meta_codes = meta.groupby(["stage", "condition", "base_task_id"], dropna=False).agg(
        observed_meta_codes=("choice_code", lambda values: ";".join(sorted(set(filter(None, map(clean, values)))))),
        observed_meta_labels_zh=("choice_label_zh", lambda values: ";".join(sorted(set(filter(None, map(clean, values)))))),
    ).reset_index()
    meta_codes = meta_codes[meta_codes["stage"].eq("C1")]
    rows = []
    for task_record in conflict.to_dict("records"):
        task, condition = clean(task_record.get("base_task_id")), normalise_condition(task_record.get("condition"))
        subset = cluster[(cluster["base_task_id"].eq(task)) & (cluster["condition"].eq(condition))].sort_values("cluster_rank")
        largest = subset[subset["cluster_rank"].astype(float).eq(1)]
        best_rank = pd.to_numeric(pd.Series([task_record.get("best_cluster_rank")]), errors="coerce").iloc[0]
        best = subset[pd.to_numeric(subset["cluster_rank"], errors="coerce").eq(best_rank)] if pd.notna(best_rank) else pd.DataFrame()
        comparison = best if not best.empty and (best_rank != 1) else subset[pd.to_numeric(subset["cluster_rank"], errors="coerce").gt(1)].head(1)
        largest_record = largest.iloc[0].to_dict() if not largest.empty else {}
        compare_record = comparison.iloc[0].to_dict() if not comparison.empty else {}
        n_pair_diff = abs(float(largest_record.get("cluster_n_pairs_median", np.nan)) - float(compare_record.get("cluster_n_pairs_median", np.nan))) if pd.notna(largest_record.get("cluster_n_pairs_median")) and pd.notna(compare_record.get("cluster_n_pairs_median")) else np.nan
        top_diff = abs(float(largest_record.get("cluster_top_y_median", np.nan)) - float(compare_record.get("cluster_top_y_median", np.nan))) if pd.notna(largest_record.get("cluster_top_y_median")) and pd.notna(compare_record.get("cluster_top_y_median")) else np.nan
        bottom_diff = abs(float(largest_record.get("cluster_bottom_y_median", np.nan)) - float(compare_record.get("cluster_bottom_y_median", np.nan))) if pd.notna(largest_record.get("cluster_bottom_y_median")) and pd.notna(compare_record.get("cluster_bottom_y_median")) else np.nan
        height_diff = abs(float(largest_record.get("cluster_room_height_median", np.nan)) - float(compare_record.get("cluster_room_height_median", np.nan))) if pd.notna(largest_record.get("cluster_room_height_median")) and pd.notna(compare_record.get("cluster_room_height_median")) else np.nan
        seam_diff = abs(float(largest_record.get("cluster_seam_pair_count_median", np.nan)) - float(compare_record.get("cluster_seam_pair_count_median", np.nan))) if pd.notna(largest_record.get("cluster_seam_pair_count_median")) and pd.notna(compare_record.get("cluster_seam_pair_count_median")) else np.nan
        x_diff = wall_x_distance(clean(largest_record.get("cluster_wall_x_median_json")), clean(compare_record.get("cluster_wall_x_median_json")))
        meta_row = meta_codes[(meta_codes["condition"].eq(condition)) & (meta_codes["base_task_id"].eq(task))]
        codes = clean(meta_row.iloc[0]["observed_meta_codes"]) if not meta_row.empty else ""
        causes: list[str] = []
        if pd.notna(n_pair_diff) and n_pair_diff >= 1:
            causes.append("different_wall_count_or_topology")
        if any(token in codes for token in ("overextension", "underextension", "oos_", "topology_issue")):
            causes.append("scope_or_adjacent_space_extent")
        if any(token in codes for token in ("seam", "seam_issue")) or (pd.notna(seam_diff) and seam_diff >= 1):
            causes.append("panorama_seam_boundary")
        if pd.notna(top_diff) and top_diff >= 10 and (pd.isna(bottom_diff) or top_diff >= bottom_diff):
            causes.append("ceiling_boundary_interpretation")
        if pd.notna(bottom_diff) and bottom_diff >= 10 and (pd.isna(top_diff) or bottom_diff > top_diff):
            causes.append("floor_boundary_interpretation")
        if pd.notna(height_diff) and height_diff >= 15:
            causes.append("room_height_or_vertical_extent")
        if x_diff is not None and x_diff >= 15:
            causes.append("horizontal_wall_boundary_placement")
        if any(token in codes for token in ("occlusion", "reflection_transparency", "low_texture", "low_image_quality")):
            causes.append("visibility_or_appearance_evidence")
        if not causes:
            causes.append("mixed_or_not_identifiable_from_available_geometry")
        priority = clean(task_record.get("crowd_gt_relationship")) not in {
            "unimodal_high_gt_alignment", "unimodal_intermediate_gt_alignment",
            "dominant_with_singleton_dissent_largest_cluster_best_gt_alignment",
        }
        rows.append({
            **task_record,
            "comparison_cluster_rank": compare_record.get("cluster_rank", np.nan),
            "comparison_cluster_support": compare_record.get("cluster_support", np.nan),
            "largest_cluster_n_pairs_median": largest_record.get("cluster_n_pairs_median", np.nan),
            "comparison_cluster_n_pairs_median": compare_record.get("cluster_n_pairs_median", np.nan),
            "n_pair_difference": n_pair_diff, "top_boundary_difference_px": top_diff,
            "bottom_boundary_difference_px": bottom_diff, "room_height_difference_px": height_diff,
            "wall_x_circular_rmse_px": x_diff, "seam_pair_count_difference": seam_diff,
            "observed_meta_codes": codes,
            "observable_geometric_difference_codes": ";".join(dict.fromkeys(causes)),
            "manual_review_priority": priority,
            "image_reference": image_map.get(task, ""),
            "interpretation_boundary_zh": "这是最大簇与最佳 GT 对齐簇（或次簇）的可观察几何差异，不把算法归类直接当作真实因果原因。",
        })
    result = pd.DataFrame(rows)
    exploded = []
    for record in result.to_dict("records"):
        for cause in clean(record["observable_geometric_difference_codes"]).split(";"):
            exploded.append({"cause_code": cause, **record})
    cause_summary = pd.DataFrame(exploded).groupby("cause_code").agg(
        task_condition_count=("base_task_id", "size"), unique_task_count=("base_task_id", "nunique"),
        manual_review_priority_count=("manual_review_priority", "sum"),
        median_GT_gap=("crowd_gt_gap_best_minus_largest", lambda values: numeric(values).median()),
    ).reset_index().sort_values("task_condition_count", ascending=False)
    return result.sort_values(["manual_review_priority", "crowd_gt_gap_best_minus_largest"], ascending=[False, False]), cause_summary


def exclusion_peer_analysis(data: Mapping[str, pd.DataFrame], flagged: pd.DataFrame) -> tuple[pd.DataFrame, pd.DataFrame]:
    peer = data["excluded_peer"].copy()
    if peer.empty:
        return pd.DataFrame(), pd.DataFrame()
    flags = flagged[flagged["stage"].eq("C1")][["base_task_id", "condition", "worker_id", "orthogonal_flags", "worker_process_class", "primary_exclusion_class"]].drop_duplicates()
    peer = peer.rename(columns={"excluded_worker_id": "worker_id"})
    peer["worker_id"] = peer["worker_id"].map(worker_id)
    peer = peer.merge(flags, how="left", on=["base_task_id", "condition", "worker_id"])
    for column in ("q_boundary", "q_wallwall"):
        peer[column] = numeric(peer[column])
    peer["q_min"] = peer[["q_boundary", "q_wallwall"]].min(axis=1)
    summary_rows = []
    for reason in sorted({token for value in peer["orthogonal_flags"].fillna("") for token in value.split(";") if token}):
        subset = peer[peer["orthogonal_flags"].fillna("").str.split(";").map(lambda values: reason in values)]
        summary_rows.append({
            "reason_code": reason, "peer_pair_count": len(subset),
            "excluded_worker_count": subset["worker_id"].nunique(), "task_count": subset["base_task_id"].nunique(),
            "q_boundary_median": float(subset["q_boundary"].median()) if subset["q_boundary"].notna().any() else np.nan,
            "q_wallwall_median": float(subset["q_wallwall"].median()) if subset["q_wallwall"].notna().any() else np.nan,
            "q_min_median": float(subset["q_min"].median()) if subset["q_min"].notna().any() else np.nan,
            "q095_compatible_rate": float(subset.get("passes_q095", pd.Series([False] * len(subset))).map(truth).mean()) if len(subset) else np.nan,
            "interpretation": "peer comparison sensitivity; does not restore formal eligibility",
        })
    return peer, pd.DataFrame(summary_rows).sort_values("peer_pair_count", ascending=False)


def worker_viewpoint_and_quality(data: Mapping[str, pd.DataFrame], flagged: pd.DataFrame) -> pd.DataFrame:
    view = data["worker_view"].copy()
    c1 = flagged[flagged["stage"].eq("C1")].copy()
    c1["iou_to_gt"] = numeric(c1["iou_to_gt"])
    c1["active_time_observed_seconds"] = numeric(c1["active_time_observed_seconds"])
    worker = c1.groupby("worker_id", as_index=False).agg(
        c1_record_count=("base_task_id", "size"), c1_task_count=("base_task_id", "nunique"),
        iou_count=("iou_to_gt", "count"), iou_mean=("iou_to_gt", "mean"), iou_median=("iou_to_gt", "median"),
        active_time_count=("active_time_observed_seconds", "count"), active_time_median=("active_time_observed_seconds", "median"),
        administrative_record_count=("flag__administrative_exclusion", "sum"),
        outside_assignment_record_count=("flag__outside_assignment", "sum"),
    )
    if not view.empty:
        view["worker_id"] = view["worker_id"].map(worker_id)
        worker = worker.merge(view, how="outer", on="worker_id")
    worker["interpretation_boundary_zh"] = "模式倾向和任务中心化角点数只能说明重复几何策略，不自动等于质量、专长或正确观点。"
    return worker.sort_values("c1_task_count", ascending=False)


def image_file_index() -> dict[str, list[Path]]:
    index: dict[str, list[Path]] = defaultdict(list)
    skip_parts = {".git", "node_modules", "venv", ".venv", "full_uncertainty_data_mining_20260821_v2", "full_uncertainty_data_mining_20260821_v3"}
    for path in ROOT.rglob("*"):
        if not path.is_file() or path.suffix.lower() not in IMAGE_SUFFIXES:
            continue
        if any(part in skip_parts for part in path.parts):
            continue
        index[path.stem].append(path)
        index[path.name].append(path)
    return index


def resolve_images(tasks: Iterable[str], image_map: Mapping[str, str]) -> pd.DataFrame:
    file_index = image_file_index()
    rows = []
    for task in sorted(set(filter(None, map(clean, tasks)))):
        reference = clean(image_map.get(task, ""))
        candidates: list[Path] = []
        ref_name = Path(reference.split("?", 1)[0]).name if reference else ""
        for key in (task, f"{task}.jpg", f"{task}.png", ref_name, Path(ref_name).stem if ref_name else ""):
            candidates.extend(file_index.get(key, []))
        candidates = list(dict.fromkeys(candidates))
        chosen = candidates[0] if candidates else None
        rows.append({
            "base_task_id": task, "building_id": building_from_task(task),
            "image_reference": reference, "repository_image_found": chosen is not None,
            "repository_image_path": chosen.relative_to(ROOT).as_posix() if chosen else "",
            "candidate_image_count": len(candidates),
        })
    return pd.DataFrame(rows)


def make_contact_sheets(out: Path, categories: Mapping[str, Sequence[str]], image_index: pd.DataFrame) -> pd.DataFrame:
    gallery_dir = out / "case_galleries"
    gallery_dir.mkdir(parents=True, exist_ok=True)
    path_map = image_index.set_index("base_task_id")["repository_image_path"].to_dict() if not image_index.empty else {}
    rows = []
    font = ImageFont.load_default()
    for category, tasks in categories.items():
        unique = list(dict.fromkeys(filter(None, map(clean, tasks))))
        page = 0
        found_count = 0
        for start in range(0, len(unique), 12):
            chunk = unique[start:start + 12]
            canvas = Image.new("RGB", (1280, 720), "white")
            draw = ImageDraw.Draw(canvas)
            page_items = 0
            for offset, task in enumerate(chunk):
                relative = clean(path_map.get(task, ""))
                if not relative or not (ROOT / relative).is_file():
                    continue
                try:
                    image = Image.open(ROOT / relative).convert("RGB")
                    image.thumbnail((300, 250))
                except Exception:
                    continue
                col, row = offset % 4, offset // 4
                x, y = 10 + col * 318, 10 + row * 235
                canvas.paste(image, (x, y))
                draw.text((x, y + 185), task[:43], fill="black", font=font)
                page_items += 1
                found_count += 1
            if not page_items:
                continue
            page += 1
            target = gallery_dir / f"{re.sub(r'[^A-Za-z0-9_-]+', '_', category)}_{page:03d}.png"
            canvas.save(target, quality=90)
            rows.append({
                "category": category, "page": page, "gallery_path": target.relative_to(out).as_posix(),
                "tasks_requested_on_page": len(chunk), "images_rendered_on_page": page_items,
            })
        if not page:
            rows.append({"category": category, "page": 0, "gallery_path": "", "tasks_requested_on_page": len(unique), "images_rendered_on_page": 0})
    return pd.DataFrame(rows)


def chart_outputs(out: Path, frames: Mapping[str, pd.DataFrame]) -> list[Path]:
    charts: list[Path] = []
    convergence = frames["semi_convergence"]
    if not convergence.empty:
        plot = convergence.sort_values("delta_shannon_entropy")
        fig, ax = plt.subplots(figsize=(12, 7))
        ax.barh(range(len(plot)), plot["delta_shannon_entropy"])
        ax.set_yticks(range(len(plot)))
        ax.set_yticklabels(plot["base_task_id"].str.slice(0, 28), fontsize=6)
        ax.axvline(0, linewidth=1)
        ax.set_xlabel("Shannon entropy difference: Semi - Manual")
        ax.set_title("All 25 paired images: Semi convergence versus expansion")
        fig.tight_layout()
        path = out / "SEMI_25_CONVERGENCE_EXPANSION.png"
        fig.savefig(path, dpi=180)
        plt.close(fig)
        charts.append(path)
    proposal = frames["proposal_task"]
    if not proposal.empty:
        fig, ax = plt.subplots(figsize=(8, 6))
        ax.scatter(proposal["edit_rate"], proposal["delta_shannon_entropy"])
        for row in proposal.itertuples(index=False):
            ax.annotate(str(row.base_task_id)[:10], (row.edit_rate, row.delta_shannon_entropy), fontsize=6)
        ax.axhline(0, linewidth=1)
        ax.set_xlabel("Worker edit rate")
        ax.set_ylabel("Semi - Manual entropy")
        ax.set_title("Proposal retention/revision and task-level uncertainty")
        fig.tight_layout()
        path = out / "PROPOSAL_EDIT_RATE_VS_UNCERTAINTY.png"
        fig.savefig(path, dpi=180)
        plt.close(fig)
        charts.append(path)
    for key, title, file_name, value_col in (
        ("exclusion_summary", "Orthogonal exclusion-reason record counts", "EXCLUSION_REASON_COUNTS.png", "record_count"),
        ("cause_summary", "Observable geometry differences in crowd/GT audit", "CROWD_GT_CAUSE_COUNTS.png", "task_condition_count"),
        ("tag_behavior_summary", "Tag-behavior audit cases", "TAG_BEHAVIOR_CASE_COUNTS.png", "row_count"),
    ):
        frame = frames[key]
        if frame.empty:
            continue
        label_col = "reason_code" if key == "exclusion_summary" else "cause_code" if key == "cause_summary" else "case_code"
        plot = frame.sort_values(value_col).tail(20)
        fig, ax = plt.subplots(figsize=(10, 6))
        ax.barh(plot[label_col], plot[value_col])
        ax.set_title(title)
        fig.tight_layout()
        path = out / file_name
        fig.savefig(path, dpi=180)
        plt.close(fig)
        charts.append(path)
    dual = frames["dual_annotator"]
    if not dual.empty and numeric(dual["cyclic_rmse_diagonal_normalized"]).notna().any():
        fig, ax = plt.subplots(figsize=(8, 5))
        ax.hist(numeric(dual["cyclic_rmse_diagonal_normalized"]).dropna(), bins=20)
        ax.set_xlabel("Cyclic coordinate RMSE / panorama diagonal")
        ax.set_ylabel("Two-annotator tasks")
        ax.set_title("Two-annotator geometry sensitivity")
        fig.tight_layout()
        path = out / "DUAL_ANNOTATOR_RMSE.png"
        fig.savefig(path, dpi=180)
        plt.close(fig)
        charts.append(path)
    return charts


def dataframe_markdown(frame: pd.DataFrame, columns: Sequence[str] | None = None, limit: int | None = None) -> str:
    if frame is None or frame.empty:
        return "无可评价记录。"
    selected = frame.copy()
    if columns:
        selected = selected[[column for column in columns if column in selected]]
    if limit:
        selected = selected.head(limit)
    return selected.to_markdown(index=False)


def build_markdown_report(out: Path, frames: Mapping[str, pd.DataFrame], validation: Mapping[str, Any]) -> str:
    convergence = frames["semi_convergence"]
    convergence_counts = convergence["uncertainty_direction"].value_counts().rename_axis("direction").reset_index(name="task_count")
    proposal_stage = frames["proposal_stage"]
    gt_rel = frames["crowd_gt_causes"]["crowd_gt_relationship"].value_counts().rename_axis("relationship").reset_index(name="task_condition_count")
    lines = [
        "# 360° 全景布局标注：完整统计、不确定性与 Proposal 行为审计（v3）",
        "",
        "## 1. 数据保留与分析边界",
        "",
        "本报告不以工人是否进入后续阶段作为删除条件。行政退出、任务外分配、旧正式资格不足、Scope 不满足、独立性不足等记录均在数据挖掘层保留；每个结局只因自身不可计算而缺失。排除原因是正交字段，因此同一行可以同时属于多个原因。",
        "",
        "模型 proposal 后的 GT/效用指标下降统一记为 `negative_metric_change`（指标负向变化）。该名称只陈述计算结果，不判断修改在真实几何上不合理。Operational GT 只有有限覆盖，而且 Manhattan 表示可能迫使场景进入特定拓扑。",
        "",
        "## 2. 数据覆盖与正交原因",
        "",
        dataframe_markdown(frames["coverage"]),
        "",
        dataframe_markdown(frames["exclusion_summary"]),
        "",
        "完整逐行标志见 `ALL_RECORDS_WITH_ORTHOGONAL_FLAGS.csv`，原因重叠见 `EXCLUSION_REASON_OVERLAPS.csv`，原因存在/不存在的结果描述见 `EXCLUSION_REASON_OUTCOME_SUMMARY.csv`。这些子组不是随机形成，不能把均值差解释为原因效应。",
        "",
        "## 3. 全部 25 张 Manual/Semi 同图任务：收敛与扩散",
        "",
        dataframe_markdown(convergence_counts),
        "",
        dataframe_markdown(convergence, [
            "base_task_id", "building_id", "image_reference", "manual_shannon_entropy", "semi_shannon_entropy",
            "delta_shannon_entropy", "uncertainty_direction", "magnitude_band", "mode_transition",
            "supported_multimodality_transition", "delta_quality_iou", "quality_direction",
        ]),
        "",
        "`semi_convergence` 只表示 Semi 的工人间模式熵低于 Manual；`semi_expansion` 表示更高。两者均不直接回答正确性。所有 25 张图片均列在上表和工作簿 `Semi25`。",
        "",
        "## 4. Proposal：结构发现、共享锚定与强制拟合的可观察证据",
        "",
        dataframe_markdown(proposal_stage),
        "",
        dataframe_markdown(frames["proposal_task"], [
            "base_task_id", "image_reference", "delta_shannon_entropy", "delta_quality_iou", "edit_rate",
            "edit_rmse_median", "positive_metric_change_count", "negative_metric_change_count",
            "proposal_response_pattern", "anchor_evidence", "structure_discovery_evidence", "forced_fit_tradeoff_evidence",
        ]),
        "",
        "可支持的判断只到以下层级：低编辑率且模式收敛与 proposal 保留相符；多人编辑后仍收敛与共享修正或锚定均相容；多人向不同方向修改并扩散与 proposal 被拒绝或图像多解相容。仅凭 GT 小幅下降，不能把修改定性为不合理。",
        "",
        "## 5. ‘模型良好/可接受’标签后仍修改角点",
        "",
        dataframe_markdown(frames["tag_behavior_summary"]),
        "",
        dataframe_markdown(frames["tag_behavior_cases"], [
            "stage", "base_task_id", "image_reference", "worker_id", "tag_codes", "case_codes",
            "edited_bool", "edit_magnitude_px_equivalent", "edit_magnitude_band", "topology_changed",
            "iou_to_gt", "cluster_rank", "is_largest_mode", "active_time_observed_seconds",
            "metric_change_class", "uncertainty_interpretation_zh",
        ]),
        "",
        "可接受标签后的任意角点移动说明仍存在残余调整，但不必然是强标注不确定性。小幅、同拓扑、仍处于主模式且 GT 对齐较高的变化更符合局部精度修正；较大坐标变化、拓扑变化、最终落入非主模式或明显偏离 GT 才构成更强的标签—几何不一致证据。",
        "",
        "## 6. ‘非常简单’标签但质量或一致性不佳",
        "",
        "`trivial_tag_but_low_GT_alignment`、`trivial_tag_but_final_nonlargest_mode`、`trivial_tag_but_structural_failure_flag`、`trivial_tag_but_long_active_time` 和 `trivial_tag_but_material_edit` 均在上一节完整列出。简单标签是工人的主观难度判断，不应被当作质量标签。",
        "",
        "## 7. 101 个 Crowd–GT 任务条件及可观察几何差异",
        "",
        dataframe_markdown(gt_rel),
        "",
        dataframe_markdown(frames["cause_summary"]),
        "",
        dataframe_markdown(frames["crowd_gt_causes"], [
            "base_task_id", "image_reference", "condition", "crowd_gt_relationship", "largest_cluster_support",
            "second_cluster_support", "best_cluster_rank", "best_cluster_support", "largest_cluster_median_iou",
            "best_cluster_median_iou", "crowd_gt_gap_best_minus_largest", "observable_geometric_difference_codes",
            "n_pair_difference", "top_boundary_difference_px", "bottom_boundary_difference_px",
            "room_height_difference_px", "wall_x_circular_rmse_px", "observed_meta_codes", "manual_review_priority",
        ]),
        "",
        "几何原因表比较最大簇与最佳 GT 对齐簇；若二者相同，则比较次簇。结果描述可观察差异：墙/角点数量、顶边、底边、房间高度、水平墙边界、全景接缝、Scope/相邻空间和可见性。它不把规则归类直接当成真实因果。",
        "",
        "## 8. 只有两名标注者的任务：无 GT 敏感性",
        "",
        dataframe_markdown(frames["dual_summary"]),
        "",
        dataframe_markdown(frames["dual_annotator"], [
            "stage", "condition", "base_task_id", "image_reference", "worker_id_left", "worker_id_right",
            "topology_agreement_class", "n_corners_left", "n_corners_right", "cyclic_rmse_diagonal_normalized",
            "coordinate_disagreement_band", "sensitivity_role",
        ]),
        "",
        "双人任务不能估计多数共识，也不能替代 GT；但可以提供拓扑是否一致、同拓扑坐标距离和阶段/模式差异的敏感性证据。",
        "",
        "## 9. 任务外、非 canonical 与修订版本",
        "",
        dataframe_markdown(frames["out_of_task_summary"]),
        "",
        dataframe_markdown(frames["out_of_task_rows"], [
            "stage", "condition", "base_task_id", "image_reference", "worker_id", "annotation_id",
            "canonical_join_status", "audit_class", "lead_time_seconds",
        ]),
        "",
        "## 10. 被排除工人与标准同任务工人的几何关系",
        "",
        dataframe_markdown(frames["excluded_peer_summary"]),
        "",
        "该表只回答这些记录在同任务中与标准工人的几何相似程度；不恢复正式资格，也不把不一致解释为工人能力。",
        "",
        "## 11. 工人重复观点、质量和时间",
        "",
        dataframe_markdown(frames["worker_integrated"], [
            "worker_id", "c1_record_count", "c1_task_count", "iou_count", "iou_mean", "active_time_count",
            "active_time_median", "largest_mode_rate", "supported_minority_mode_rate", "mean_task_centered_n_pairs",
            "administrative_record_count", "outside_assignment_record_count",
        ]),
        "",
        "## 12. 完整图片实例索引",
        "",
        dataframe_markdown(frames["image_index"]),
        "",
        "每类分析的全部实例都通过 `base_task_id` 和 `image_reference` 列出。仓库中可定位到原图时，`case_galleries/` 同时生成缩略图分页；找不到原图不会删除该实例。",
        "",
        "## 13. 统计和解释限制",
        "",
        "1. Manual/Semi 不是完整图像级随机试验，任务和工人组成不同。",
        "2. GT 覆盖为 101 个 task-condition，且 operational GT 不是不可质疑的真实世界真值。",
        "3. 两名标注者只能形成成对敏感性，不能形成稳定多数模式。",
        "4. Proposal 收敛可能来自正确结构提示，也可能来自共享初始化锚定或 Manhattan 强制拟合。",
        "5. 元标签是主观判断；`acceptable` 和 `trivial` 均不能替代几何、GT、共识和时间证据。",
        "6. 排除原因高度重叠，原因子组的均值差不能解释为排除原因的因果效应。",
        "",
        "## 14. 可复现性",
        "",
        f"生成时间：{datetime.now(timezone.utc).isoformat()}。",
        f"校验摘要：`{json.dumps(validation, ensure_ascii=False, sort_keys=True)}`。",
        "",
        "变量中文解释和近似计算见 `DATA_DICTIONARY_ZH_V3.csv`；输入/输出 SHA-256 见 `INPUT_PROVENANCE_V3.csv` 和 `OUTPUT_MANIFEST.csv`。",
    ]
    report = "\n".join(lines) + "\n"
    forbidden = ["harmful_rate", "有害修改", "过度修正", "over-correction"]
    for token in forbidden:
        if token in report:
            raise AssertionError(f"forbidden causal terminology leaked into report: {token}")
    return report


def docx_set_cell_text(cell, value: Any, font_size: float = 7.5) -> None:
    cell.text = "" if pd.isna(value) else str(value)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(font_size)
            run.font.name = "Arial"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def add_docx_table(document: Document, frame: pd.DataFrame, columns: Sequence[str], title: str, max_rows: int | None = None) -> None:
    document.add_heading(title, level=2)
    if frame.empty:
        document.add_paragraph("无可评价记录。")
        return
    selected = frame[[column for column in columns if column in frame]].copy()
    if max_rows is not None:
        selected = selected.head(max_rows)
    table = document.add_table(rows=1, cols=len(selected.columns))
    table.style = "Table Grid"
    for index, column in enumerate(selected.columns):
        docx_set_cell_text(table.rows[0].cells[index], column, 7)
        table.rows[0].cells[index]._tc.get_or_add_tcPr().append(OxmlElement("w:shd"))
    for record in selected.itertuples(index=False):
        cells = table.add_row().cells
        for index, value in enumerate(record):
            if isinstance(value, float) and math.isfinite(value):
                value = f"{value:.5g}"
            docx_set_cell_text(cells[index], value)
    document.add_paragraph(f"记录数：{len(frame)}；本表显示：{len(selected)}。")


def build_docx(out: Path, report_path: Path, frames: Mapping[str, pd.DataFrame], charts: Sequence[Path], galleries: pd.DataFrame) -> Path:
    target = out / "Paper_A_完整数据整理与不确定性分析报告_20260821_v3.docx"
    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(1.6)
    section.bottom_margin = Cm(1.6)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    styles["Normal"].font.size = Pt(9)
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("360° 全景布局标注\n完整统计、不确定性与 Proposal 行为审计（v3）")
    run.bold = True
    run.font.size = Pt(18)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    document.add_paragraph("客观数据整理版本。所有排除记录在可计算范围内保留；GT/指标负向变化采用中性命名。")

    document.add_heading("核心统计", level=1)
    add_docx_table(document, frames["coverage"], list(frames["coverage"].columns), "数据覆盖", 20)
    add_docx_table(document, frames["exclusion_summary"], ["reason_code", "reason_zh", "record_count", "worker_count", "task_count", "geometry_computable_count", "quality_computable_count", "active_time_computable_count"], "正交排除原因", None)
    add_docx_table(document, frames["semi_convergence"], ["base_task_id", "image_reference", "delta_shannon_entropy", "uncertainty_direction", "mode_transition", "delta_quality_iou", "quality_direction"], "全部 25 张 Semi 收敛/扩散图片", None)
    add_docx_table(document, frames["proposal_task"], ["base_task_id", "image_reference", "delta_shannon_entropy", "edit_rate", "edit_rmse_median", "positive_metric_change_count", "negative_metric_change_count", "proposal_response_pattern", "anchor_evidence", "structure_discovery_evidence", "forced_fit_tradeoff_evidence"], "Proposal 任务级审计", None)
    add_docx_table(document, frames["tag_behavior_cases"], ["stage", "base_task_id", "image_reference", "worker_id", "tag_codes", "case_codes", "edited_bool", "edit_magnitude_px_equivalent", "topology_changed", "iou_to_gt", "cluster_rank", "metric_change_class", "uncertainty_interpretation_zh"], "标签—行为完整实例", None)

    document.add_page_break()
    document.add_heading("Crowd、少数模式与 GT", level=1)
    add_docx_table(document, frames["cause_summary"], list(frames["cause_summary"].columns), "可观察几何差异分类", None)
    add_docx_table(document, frames["crowd_gt_causes"], ["base_task_id", "image_reference", "condition", "crowd_gt_relationship", "largest_cluster_support", "second_cluster_support", "best_cluster_rank", "best_cluster_support", "crowd_gt_gap_best_minus_largest", "observable_geometric_difference_codes", "n_pair_difference", "top_boundary_difference_px", "bottom_boundary_difference_px", "wall_x_circular_rmse_px", "manual_review_priority"], "全部 101 个 Crowd–GT task-condition", None)

    document.add_page_break()
    document.add_heading("双人敏感性、任务外与排除数据", level=1)
    add_docx_table(document, frames["dual_summary"], list(frames["dual_summary"].columns), "双人任务阶段汇总", None)
    add_docx_table(document, frames["dual_annotator"], ["stage", "condition", "base_task_id", "image_reference", "worker_id_left", "worker_id_right", "topology_agreement_class", "n_corners_left", "n_corners_right", "cyclic_rmse_diagonal_normalized", "coordinate_disagreement_band"], "所有双人任务", None)
    add_docx_table(document, frames["out_of_task_rows"], ["stage", "condition", "base_task_id", "image_reference", "worker_id", "annotation_id", "canonical_join_status", "audit_class", "lead_time_seconds"], "任务外/非 canonical/修订版本", None)
    add_docx_table(document, frames["excluded_peer_summary"], list(frames["excluded_peer_summary"].columns), "被排除记录的同任务 Peer 敏感性", None)

    document.add_page_break()
    document.add_heading("工人、时间和图片索引", level=1)
    add_docx_table(document, frames["worker_integrated"], ["worker_id", "c1_record_count", "c1_task_count", "iou_count", "iou_mean", "active_time_count", "active_time_median", "largest_mode_rate", "supported_minority_mode_rate", "mean_task_centered_n_pairs", "administrative_record_count", "outside_assignment_record_count"], "工人综合描述", None)
    add_docx_table(document, frames["image_index"], list(frames["image_index"].columns), "完整图片实例索引", None)

    document.add_page_break()
    document.add_heading("图表与缩略图", level=1)
    for path in charts:
        if path.is_file():
            document.add_paragraph(path.stem)
            document.add_picture(str(path), width=Inches(6.8))
    if not galleries.empty:
        for record in galleries.itertuples(index=False):
            if record.gallery_path and (out / record.gallery_path).is_file():
                document.add_paragraph(f"{record.category} — page {record.page}")
                document.add_picture(str(out / record.gallery_path), width=Inches(6.8))

    document.add_page_break()
    document.add_heading("变量与计算", level=1)
    add_docx_table(document, frames["dictionary"], ["variable_en", "variable_zh", "meaning_zh", "approximate_formula"], "中英文数据词典", None)
    document.save(target)
    opened = Document(target)
    if len(opened.paragraphs) < 20 or len(opened.tables) < 8:
        raise AssertionError("DOCX validation failed")
    return target


def safe_sheet_name(name: str, existing: set[str]) -> str:
    base = re.sub(r"[\\/*?:\[\]]", "_", name)[:31] or "Sheet"
    candidate, index = base, 2
    while candidate in existing:
        suffix = f"_{index}"
        candidate = base[:31 - len(suffix)] + suffix
        index += 1
    existing.add(candidate)
    return candidate


def build_xlsx(out: Path, frames: Mapping[str, pd.DataFrame], charts: Sequence[Path]) -> Path:
    target = out / "Paper_A_完整数据整理与分析工作簿_20260821_v3.xlsx"
    workbook = Workbook()
    workbook.remove(workbook.active)
    sheet_map = [
        ("README_说明", frames["readme"]),
        ("Coverage_覆盖", frames["coverage"]),
        ("Semi25_收敛扩散", frames["semi_convergence"]),
        ("Proposal_任务", frames["proposal_task"]),
        ("Proposal_阶段", frames["proposal_stage"]),
        ("TagBehavior_全行", frames["tag_behavior_all"]),
        ("TagBehavior_案例", frames["tag_behavior_cases"]),
        ("TagBehavior_汇总", frames["tag_behavior_summary"]),
        ("GT101_几何原因", frames["crowd_gt_causes"]),
        ("GT_原因汇总", frames["cause_summary"]),
        ("GT_簇明细", frames["crowd_cluster"]),
        ("Dual2_双人", frames["dual_annotator"]),
        ("Dual2_汇总", frames["dual_summary"]),
        ("Exclusion_逐行", frames["flagged_records"]),
        ("Exclusion_汇总", frames["exclusion_summary"]),
        ("Exclusion_重叠", frames["exclusion_overlap"]),
        ("Exclusion_结果", frames["exclusion_outcomes"]),
        ("Exclusion_Peer", frames["excluded_peer"]),
        ("OutOfTask_逐行", frames["out_of_task_rows"]),
        ("OutOfTask_汇总", frames["out_of_task_summary"]),
        ("Worker_综合", frames["worker_integrated"]),
        ("Image_图片索引", frames["image_index"]),
        ("Gallery_缩略图", frames["galleries"]),
        ("Dictionary_词典", frames["dictionary"]),
        ("Inputs_输入", frames["input_provenance"]),
    ]
    existing: set[str] = set()
    for desired, frame in sheet_map:
        name = safe_sheet_name(desired, existing)
        ws = workbook.create_sheet(name)
        if frame is None or frame.empty:
            ws.cell(1, 1, "无可评价记录")
            continue
        clean_frame = frame.copy()
        for column in clean_frame.columns:
            clean_frame[column] = clean_frame[column].map(
                lambda value: json.dumps(value, ensure_ascii=False, default=str) if isinstance(value, (dict, list, tuple, set)) else value
            )
        ws.append(list(clean_frame.columns))
        for record in clean_frame.itertuples(index=False, name=None):
            ws.append([None if pd.isna(value) else value for value in record])
        ws.freeze_panes = "A2"
        ws.auto_filter.ref = ws.dimensions
        header_fill = PatternFill("solid", fgColor="D9EAF7")
        for cell in ws[1]:
            cell.font = Font(bold=True)
            cell.fill = header_fill
            cell.alignment = Alignment(wrap_text=True, vertical="top")
        for column_index, column in enumerate(clean_frame.columns, 1):
            values = [str(column)] + [str(value) for value in clean_frame[column].head(250).fillna("")]
            width = min(45, max(10, int(np.quantile([len(value) for value in values], 0.9)) + 2))
            ws.column_dimensions[get_column_letter(column_index)].width = width
        for row in ws.iter_rows(min_row=2):
            for cell in row:
                cell.alignment = Alignment(wrap_text=True, vertical="top")
        numeric_columns = [index + 1 for index, column in enumerate(clean_frame.columns) if pd.api.types.is_numeric_dtype(clean_frame[column])]
        for column_index in numeric_columns[:10]:
            letter = get_column_letter(column_index)
            if ws.max_row >= 3:
                ws.conditional_formatting.add(
                    f"{letter}2:{letter}{ws.max_row}",
                    ColorScaleRule(start_type="min", start_color="FEE2E2", mid_type="percentile", mid_value=50, mid_color="FFF7CC", end_type="max", end_color="DCFCE7"),
                )
    overview = workbook["README_说明"]
    row = overview.max_row + 3
    for path in charts[:4]:
        if path.is_file():
            image = XLImage(str(path))
            image.width = 560
            image.height = 330
            overview.add_image(image, f"A{row}")
            row += 18
    workbook.save(target)
    opened = load_workbook(target, read_only=True, data_only=True)
    if len(opened.sheetnames) < 20:
        raise AssertionError("XLSX validation failed")
    opened.close()
    return target


def data_dictionary_v3() -> pd.DataFrame:
    rows = [
        ("negative_metric_change", "指标负向变化", "最终 GT/效用指标比初始值低超过 0.01；不等同于真实几何修改不合理", "delta_U < -0.01"),
        ("positive_metric_change", "指标正向变化", "最终 GT/效用指标比初始值高超过 0.01", "delta_U > 0.01"),
        ("near_zero_metric_change", "近零指标变化", "变化绝对值不超过 0.01", "|delta_U| <= 0.01"),
        ("semi_convergence", "Semi 收敛", "同图等支持量下 Semi 的 Shannon 模式熵低于 Manual", "H_semi - H_manual < 0"),
        ("semi_expansion", "Semi 扩散", "同图等支持量下 Semi 的 Shannon 模式熵高于 Manual", "H_semi - H_manual > 0"),
        ("delta_shannon_entropy", "Semi−Manual 熵差", "负值表示 Semi 分布更集中，正值表示更分散", "H_semi - H_manual"),
        ("acceptable_proposal_with_any_corner_edit", "可接受 proposal 后仍编辑角点", "工人选择 acceptable 标签但几何编辑量大于零", "acceptable_tag AND edited_bool"),
        ("material_edit", "实质坐标/拓扑编辑", "编辑量位于正编辑分布中段以上，或观察到拓扑变化", "edit_band in {moderate,large} OR topology_changed"),
        ("trivial_tag_but_low_GT_alignment", "简单标签但 GT 对齐较低", "选择 trivial 且 IoU<0.8或低于同任务中位数0.1以上", "trivial AND (IoU<0.8 OR IoU<task_median-0.1)"),
        ("orthogonal_flag_count", "正交原因数量", "一条记录同时命中的排除/不可计算原因数量", "sum_j flag_j"),
        ("cyclic_rmse_diagonal_normalized", "循环对齐坐标 RMSE", "两份相同拓扑全景几何经循环/反向对齐后的坐标差，除以全景对角线", "min_alignment RMSE / sqrt(1024^2+512^2)"),
        ("observable_geometric_difference_codes", "可观察几何差异类别", "最大簇与最佳 GT 对齐簇/次簇之间的墙数、顶底边、水平边界、接缝、Scope等差异", "rule-based multi-label audit"),
        ("crowd_gt_gap_best_minus_largest", "最佳 GT 簇与最大簇差", "最佳 GT 对齐簇中位 IoU减最大簇中位 IoU", "max_m median(IoU_m)-median(IoU_largest)"),
        ("coordinate_disagreement_band", "双人坐标分歧分位带", "双人任务的循环 RMSE 在全部双人任务中的三分位位置", "tertiles of pair RMSE"),
        ("anchor_evidence", "共享锚定相容证据", "低编辑率/小幅共享修正伴随 Semi 收敛；只是相容性分类，不是因果证明", "rule-based descriptive class"),
        ("structure_discovery_evidence", "结构发现相容证据", "多人编辑、指标正向变化多于负向且未扩散时标记为相容；不证明因果", "descriptive rule"),
        ("forced_fit_tradeoff_evidence", "强制拟合/指标权衡可能性", "模式收敛同时 GT/效用指标负向或负向行多于正向行", "descriptive rule"),
    ]
    base = read_csv(V2 / "DATA_DICTIONARY_ZH.csv", required=False)
    extra = pd.DataFrame(rows, columns=["variable_en", "variable_zh", "meaning_zh", "approximate_formula"])
    return pd.concat([base, extra], ignore_index=True, sort=False).drop_duplicates("variable_en", keep="last")


def input_provenance() -> pd.DataFrame:
    paths = [
        V2 / "UNIFIED_SUBMISSION_EVIDENCE_REVIEWED.csv",
        V2 / "RAW_META_LABEL_RESPONSE_LONG.csv",
        V2 / "GEOMETRY_TASK_UNCERTAINTY_ALL_STAGES.csv",
        V2 / "GEOMETRY_PAIRWISE_DISPERSION_ALL_STAGES.csv",
        V2 / "C1_CROWD_GT_CONFLICT_TASKS.csv",
        V2 / "C1_CROWD_GT_CLUSTER_METRICS.csv",
        C1_MINING / "POPULATION_TASK_METRICS.csv",
        C1_MINING / "ROW_INCLUSION_CLASSIFICATION.csv",
        C1_MINING / "QUALITY_DATA_MINING_CONTEXTS.csv",
        C1_MINING / "EXCLUDED_WORKER_PEER_COMPARISONS.csv",
        PACKAGE / "raw_annotation_fact.csv",
        PACKAGE / "semi_review_fact.csv",
        C1_AUDIT / "c1_canonical_geometry.jsonl",
    ]
    rows = []
    for path in paths:
        rows.append({
            "path": path.relative_to(ROOT).as_posix(), "exists": path.is_file(),
            "size_bytes": path.stat().st_size if path.is_file() else np.nan,
            "sha256": sha256_file(path) if path.is_file() else "",
        })
    return pd.DataFrame(rows)


def output_manifest(out: Path) -> pd.DataFrame:
    rows = []
    for path in sorted(item for item in out.rglob("*") if item.is_file() and item.name != "OUTPUT_MANIFEST.csv"):
        rows.append({
            "path": path.relative_to(out).as_posix(), "size_bytes": path.stat().st_size,
            "sha256": sha256_file(path),
        })
    return pd.DataFrame(rows)


def create_zip(source_dir: Path, target: Path, exclude: set[Path] | None = None) -> None:
    exclude = exclude or set()
    with zipfile.ZipFile(target, "w", compression=zipfile.ZIP_DEFLATED, compresslevel=6) as archive:
        for path in sorted(item for item in source_dir.rglob("*") if item.is_file()):
            if path in exclude or path == target:
                continue
            archive.write(path, path.relative_to(source_dir))
    with zipfile.ZipFile(target, "r") as archive:
        bad = archive.testzip()
        if bad:
            raise AssertionError(f"corrupt zip member: {bad}")


def copy_analysis_code(out: Path) -> None:
    code_dir = out / "analysis_code"
    code_dir.mkdir(parents=True, exist_ok=True)
    for name in (
        "full_uncertainty_common.py", "full_uncertainty_geometry.py", "full_uncertainty_reviewed.py",
        "materialize_full_uncertainty_data_mining.py", "materialize_full_uncertainty_data_mining_v2.py",
        "run_full_uncertainty_reviewed.py", "materialize_full_uncertainty_data_mining_v3.py",
    ):
        source = ROOT / "tools" / "thesis_main" / "analysis" / name
        if source.is_file():
            shutil.copy2(source, code_dir / name)


def materialize(out: Path) -> dict[str, Any]:
    if out.exists():
        shutil.rmtree(out)
    out.mkdir(parents=True)
    data = load_data()
    image_map = build_image_reference_map(data["raw"])
    flagged, exclusion_summary, exclusion_outcomes, exclusion_overlap = orthogonal_record_flags(data["unified"])
    out_rows, out_summary = raw_out_of_task(data["raw"], flagged)
    convergence = semi_convergence_expansion(data, image_map)
    review = prepare_semi_review(data)
    proposal_stage = proposal_stage_summary(review)
    proposal_task = proposal_task_analysis(review, convergence, image_map)
    tag_all, tag_cases, tag_summary = tag_behavior_analysis(data, review, image_map)
    dual, dual_summary = dual_annotator_sensitivity(data, image_map)
    gt_causes, cause_summary = crowd_gt_geometric_causes(data, image_map)
    excluded_peer, excluded_peer_summary = exclusion_peer_analysis(data, flagged)
    worker_integrated = worker_viewpoint_and_quality(data, flagged)
    coverage = read_csv(V2 / "DATA_COVERAGE_BY_STAGE_MODE.csv")
    image_tasks = set(flagged["base_task_id"]) | set(convergence["base_task_id"]) | set(gt_causes["base_task_id"]) | set(dual["base_task_id"]) | set(tag_cases["base_task_id"])
    image_index = resolve_images(image_tasks, image_map)
    dictionary = data_dictionary_v3()
    provenance = input_provenance()
    readme = pd.DataFrame([
        {"section": "population", "description_zh": "全部可计算记录保留；旧资格和排除原因作为正交描述字段。"},
        {"section": "metric terminology", "description_zh": "GT/效用下降记为 negative_metric_change，不据此判断真实几何修改是否不合理。"},
        {"section": "Semi convergence", "description_zh": "仅表示工人间模式熵下降，不等于质量提升。"},
        {"section": "GT101", "description_zh": "101个 crowd–GT task-condition 全量列出；多数簇不自动替换 GT。"},
        {"section": "Dual annotator", "description_zh": "两人任务只作 pairwise 敏感性，不形成多数共识。"},
        {"section": "Image instances", "description_zh": "所有匹配实例均列出 base_task_id 和 image_reference；原图存在时另生成缩略图。"},
    ])
    frames: dict[str, pd.DataFrame] = {
        "readme": readme,
        "coverage": coverage,
        "flagged_records": flagged,
        "exclusion_summary": exclusion_summary,
        "exclusion_outcomes": exclusion_outcomes,
        "exclusion_overlap": exclusion_overlap,
        "out_of_task_rows": out_rows,
        "out_of_task_summary": out_summary,
        "semi_convergence": convergence,
        "proposal_stage": proposal_stage,
        "proposal_task": proposal_task,
        "tag_behavior_all": tag_all,
        "tag_behavior_cases": tag_cases,
        "tag_behavior_summary": tag_summary,
        "dual_annotator": dual,
        "dual_summary": dual_summary,
        "crowd_gt_causes": gt_causes,
        "cause_summary": cause_summary,
        "crowd_cluster": data["crowd_cluster"],
        "excluded_peer": excluded_peer,
        "excluded_peer_summary": excluded_peer_summary,
        "worker_integrated": worker_integrated,
        "image_index": image_index,
        "dictionary": dictionary,
        "input_provenance": provenance,
    }
    categories = {
        "semi_convergence_all": convergence.loc[convergence["uncertainty_direction"].eq("semi_convergence"), "base_task_id"].tolist(),
        "semi_expansion_all": convergence.loc[convergence["uncertainty_direction"].eq("semi_expansion"), "base_task_id"].tolist(),
        "acceptable_proposal_with_edit": tag_cases.loc[tag_cases["case_codes"].str.contains("acceptable_proposal_with_any_corner_edit", na=False), "base_task_id"].tolist(),
        "trivial_tag_quality_or_consensus_case": tag_cases.loc[tag_cases["case_codes"].str.contains("trivial_tag", na=False), "base_task_id"].tolist(),
        "crowd_GT_manual_review_priority": gt_causes.loc[gt_causes["manual_review_priority"], "base_task_id"].tolist(),
        "two_annotator_upper_disagreement": dual.loc[dual["coordinate_disagreement_band"].eq("upper_third"), "base_task_id"].tolist(),
    }
    galleries = make_contact_sheets(out, categories, image_index)
    frames["galleries"] = galleries
    for name, frame in frames.items():
        if name in {"readme", "crowd_cluster", "galleries"}:
            continue
        write_csv(out / f"{name.upper()}.csv", frame)
    # Canonical, descriptive filenames for navigation.
    aliases = {
        "ALL_RECORDS_WITH_ORTHOGONAL_FLAGS.csv": flagged,
        "EXCLUSION_REASON_SUMMARY.csv": exclusion_summary,
        "EXCLUSION_REASON_OUTCOME_SUMMARY.csv": exclusion_outcomes,
        "EXCLUSION_REASON_OVERLAPS.csv": exclusion_overlap,
        "OUT_OF_TASK_AND_NONSELECTED_ROWS.csv": out_rows,
        "OUT_OF_TASK_SUMMARY.csv": out_summary,
        "SEMI_25_CONVERGENCE_EXPANSION_ALL_IMAGES.csv": convergence,
        "PROPOSAL_STAGE_METRIC_CHANGE_SUMMARY.csv": proposal_stage,
        "PROPOSAL_TASK_ANCHORING_AND_DISCOVERY.csv": proposal_task,
        "TAG_BEHAVIOR_ALL_SEMI_ROWS.csv": tag_all,
        "TAG_BEHAVIOR_ALL_CASES.csv": tag_cases,
        "TAG_BEHAVIOR_CASE_SUMMARY.csv": tag_summary,
        "DUAL_ANNOTATOR_GEOMETRY_SENSITIVITY_ALL.csv": dual,
        "DUAL_ANNOTATOR_STAGE_SUMMARY.csv": dual_summary,
        "CROWD_GT_101_GEOMETRIC_CAUSE_AUDIT.csv": gt_causes,
        "CROWD_GT_GEOMETRIC_CAUSE_SUMMARY.csv": cause_summary,
        "EXCLUDED_RECORD_PEER_COMPARISONS_BY_REASON.csv": excluded_peer,
        "EXCLUDED_RECORD_PEER_SUMMARY_BY_REASON.csv": excluded_peer_summary,
        "WORKER_VIEWPOINT_QUALITY_TIME_INTEGRATED.csv": worker_integrated,
        "ALL_IMAGE_INSTANCE_INDEX.csv": image_index,
        "GALLERY_INDEX.csv": galleries,
        "DATA_DICTIONARY_ZH_V3.csv": dictionary,
        "INPUT_PROVENANCE_V3.csv": provenance,
    }
    for name, frame in aliases.items():
        write_csv(out / name, frame)
    charts = chart_outputs(out, frames)
    validation = {
        "paired_manual_semi_task_count": len(convergence),
        "crowd_gt_task_condition_count": len(gt_causes),
        "dual_annotator_task_count": len(dual),
        "unique_image_task_count": len(image_index),
        "tag_behavior_case_row_count": len(tag_cases),
        "orthogonal_flag_record_count": len(flagged),
        "negative_metric_change_terminology_only": True,
        "generated_at_utc": datetime.now(timezone.utc).isoformat(),
    }
    if validation["paired_manual_semi_task_count"] != 25:
        raise AssertionError(validation)
    if validation["crowd_gt_task_condition_count"] != 101:
        raise AssertionError(validation)
    report = build_markdown_report(out, frames, validation)
    report_path = out / "FULL_UNCERTAINTY_DATA_REPORT_ZH_V3.md"
    report_path.write_text(report, encoding="utf-8")
    docx = build_docx(out, report_path, frames, charts, galleries)
    xlsx = build_xlsx(out, frames, charts)
    copy_analysis_code(out)
    (out / "SUPERSESSION_NOTICE.md").write_text(
        "# v3 supersession notice\n\n"
        "v3 retains v2 as a reproducibility baseline but supersedes its interpretive terminology. "
        "Metric decreases are reported neutrally, and all proposal/tag conclusions are separated from causal claims.\n",
        encoding="utf-8",
    )
    (out / "REMOTE_DELIVERY_INDEX.md").write_text(
        "# Remote delivery index\n\n"
        "- `FULL_UNCERTAINTY_DATA_REPORT_ZH_V3.md` — objective report\n"
        f"- `{docx.name}` — advisor-readable DOCX\n"
        f"- `{xlsx.name}` — complete workbook\n"
        "- `Paper_A_完整数据整理报告与全部分析结果_20260821_v3.zip` — complete generated package\n"
        "- all CSV, PNG, gallery, source-code and SHA files are stored beside these files.\n",
        encoding="utf-8",
    )
    write_json(out / "VALIDATION_SUMMARY.json", validation)
    manifest = output_manifest(out)
    write_csv(out / "OUTPUT_MANIFEST.csv", manifest)
    v2_zip = out / "full_uncertainty_data_mining_20260821_v2_results_snapshot.zip"
    create_zip(V2, v2_zip)
    package = out / "Paper_A_完整数据整理报告与全部分析结果_20260821_v3.zip"
    create_zip(out, package, exclude={package})
    # Refresh external manifest after ZIP creation; package contains the previous manifest by design.
    manifest = output_manifest(out)
    write_csv(out / "OUTPUT_MANIFEST.csv", manifest)
    with zipfile.ZipFile(package, "r") as archive:
        if archive.testzip():
            raise AssertionError("v3 package zip validation failed")
    validation.update({
        "docx_size_bytes": docx.stat().st_size,
        "xlsx_size_bytes": xlsx.stat().st_size,
        "v2_snapshot_zip_size_bytes": v2_zip.stat().st_size,
        "v3_package_zip_size_bytes": package.stat().st_size,
        "output_file_count": len([path for path in out.rglob("*") if path.is_file()]),
    })
    write_json(out / "VALIDATION_SUMMARY.json", validation)
    print(json.dumps(validation, ensure_ascii=False, indent=2, sort_keys=True))
    return validation


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--output-dir", type=Path, default=DEFAULT_OUTPUT)
    args = parser.parse_args()
    materialize(args.output_dir.resolve())


if __name__ == "__main__":
    main()
